import tkinter as tk
from tkinter import ttk, messagebox
import tkinter.font as tkFont
import sqlite3
from datetime import datetime
from PIL import Image, ImageTk
import requests
import matplotlib as plot
import webbrowser
import subprocess
import os
import cv2
import mediapipe as mp
import time 
import threading
import socket
class LoginWindow:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("智心方舟——多模态数据融合的双心室心肌病智能诊断系统登录")
        self.root.geometry("1200x800")
        
        # 加载背景图片
        bg_image = Image.open(r"C:\Users\32572\Desktop\毕业论文插图\1743393571280.jpg")
        bg_image = bg_image.resize((1200, 800), Image.Resampling.LANCZOS)
        self.bg_photo = ImageTk.PhotoImage(bg_image)
    
         # 创建背景标签
        bg_label = tk.Label(self.root, image=self.bg_photo)
        bg_label.place(x=0, y=0, relwidth=1, relheight=1)
        
        # 创建半透明登录框，放置在左侧
        login_frame = tk.Frame(self.root, bg='white')
        login_frame.place(x=50, y=300, width=400, height=1800)  # 增加高度以容纳新文字
        
        ttk.Label(login_frame, text="欢迎使用多模态数据融合的双心室心肌病智能诊断系统", font=("Microsoft YaHei", 12, "bold")).pack(pady=10)
        ttk.Label(login_frame, text="保护患者隐私，维护数据安全", font=("Microsoft YaHei", 12)).pack(pady=5)
        ttk.Label(login_frame, text="请使用您的账号登录", font=("Microsoft YaHei", 10)).pack(pady=5)
     
        # 登录界面元素
        ttk.Label(login_frame, text="用户名:").pack(pady=10)
        self.username = ttk.Entry(login_frame)
        self.username.pack()
        ttk.Label(login_frame, text="密码:").pack(pady=10)
        self.password = ttk.Entry(login_frame, show="*")
        self.password.pack()
        
        ttk.Button(login_frame, text="登录", command=self.login).pack(pady=20)
        
        self.root.mainloop()
    
    def login(self):
        if self.username.get() == "202110916216" and self.password.get() == "ma200206":
            self.root.destroy()  # 关闭登录窗口
            # 启动主系统
            root = tk.Tk()
            app = HospitalManagementSystem(root)
            root.mainloop()
        else:
            messagebox.showerror("错误", "用户名或密码错误！")
#多模态数据融合系统，主要函数功能体
class HospitalManagementSystem:
    def __init__(self, root):
        import cv2
        import PIL
        from PIL import Image, ImageTk
        self.cv2 = cv2  # 保存为实例变量
        self.Image = Image
        self.ImageTk = ImageTk
        self.root = root
        self.root.title("智心方舟——基于多模态数据融合的双心室心肌病智能诊断系统V1.0")
        self.root.geometry("1200x800")
        
        # 初始化数据库连接
        self.conn = None
        self.cursor = None
        self.init_database()
        
        # 界面初始化
        self.main_frame = ttk.Frame(self.root)
        self.main_frame.pack(fill=tk.BOTH, expand=True, padx=40, pady=30)
     

        # 新增现代化配色方案
        self.primary_color = "#2E7D32"  # 主色调
        self.secondary_color = "#FFFFFF"  # 辅助色
        self.accent_color = "#81C784"   # 强调色
        self.current_patient_id = None  # 新增
        self.current_patient_name = None  # 新增
        # 配置全局样式
        style = ttk.Style()
        style.theme_use("clam")
        # 自定义按钮样式
        style.configure("Custom.TButton", 
                   padding=12,
                   font=("Microsoft YaHei", 12, "bold"),
                   foreground=self.secondary_color,
                   background=self.primary_color,
                   bordercolor=self.accent_color,
                   relief="flat")
        style.map("Custom.TButton",
            background=[("active", self.accent_color), ("disabled", "#BDBDBD")])
    
        # 表格样式优化
        style.configure("Treeview",
                  font=("Microsoft YaHei", 10),
                  rowheight=30,
                  fieldbackground=self.secondary_color)
        style.configure("Treeview.Heading",
                  font=("Microsoft YaHei", 11, "bold"),
                  background=self.primary_color,
                  foreground="white")
        style.map("Treeview",
            background=[('selected', self.accent_color)],
            foreground=[('selected', 'white')])
    
        # 输入框样式
        style.configure("TCustom.TEntry",
                  bordercolor=self.primary_color,
                  lightcolor=self.accent_color,
                  darkcolor=self.primary_color)
        style.map("TCustom.TEntry",
            bordercolor=[('active', self.primary_color)])
        # 系统标题
        title_font = tkFont.Font(family="Microsoft YaHei", size=20, weight="bold")
        ttk.Label(
            self.main_frame,
            text="欢迎使用多模态数据融合的双心室心肌病智能诊断系统V1.0",
            font=title_font,
            wraplength=1000
        ).pack(pady=(0, 30))
        # 功能按钮
        button_frame = ttk.Frame(self.main_frame)
        button_frame.pack(fill=tk.BOTH, expand=True, padx=50)
        
        buttons = [
            ("患者信息管理", "管理患者信息"), 
            ("医生信息管理", "管理医生信息"),
            ("病例信息管理(XGboost)", "病例信息管理"),
            ("心脏病药品管理", "智能药品管理"),
            ("心率血氧（单片机）", "心率血氧监测"),
            ("心脏影像处理（深度学习）", "智能医学影像分析"),
            ("心脏电固液耦合（计算血流动力学）", "心脏动力学仿真平台"),
            ("心脏病生化检验", "心脏生化指标检测"), 
            ("诊断终端(知识推理)", "智能诊断模块"),
            ("远程会诊中心(opencv)", "远程视频会诊"),
        ]
        # 按钮样式配置
        style = ttk.Style()
        style.configure("Custom.TButton", padding=15, font=("Microsoft YaHei", 12))
        # 创建按钮网格 - 修改为4x2布局
        for i, (text, tooltip) in enumerate(buttons):
            row = i // 2  # 每行2个按钮
            col = i % 2   # 2列
            btn = ttk.Button(
                button_frame,
                text=text,
                style="Custom.TButton",
                command=lambda t=text: self.show_message(t)
            )
            btn.grid(row=row, column=col, padx=20, pady=20, sticky="nsew")
            self.create_tooltip(btn, tooltip)
            
        # 网格布局配置 - 修改为4行2列
        for i in range(4):  # 4行
            button_frame.grid_rowconfigure(i, weight=1)
        for i in range(2):  # 2列
            button_frame.grid_columnconfigure(i, weight=1)
            
            status_frame = ttk.Frame(self.main_frame)
        status_frame.pack(fill=tk.X, side=tk.BOTTOM, pady=10)
        
        # 系统状态标签
        self.status_label = ttk.Label(status_frame,
            text="系统状态: 正常运行中",
            font=("Microsoft YaHei", 10)
        )
        self.status_label.pack(side=tk.LEFT)


          # 动态时间标签（新增）
        self.time_label = ttk.Label(
            status_frame,
            font=("Microsoft YaHei", 10)
        )
        self.time_label.pack(side=tk.LEFT, padx=20)
        self.update_clock()  # 启动时钟更新

        # 新增作者信息
        ttk.Label(
            status_frame,
            text="制作人：马文轩  指导教师：  单位：河南工程学院 @版权所有",
           
            font=("Microsoft YaHei", 10)
        ).pack(side=tk.RIGHT)
 
    # 新增时钟更新方法（添加在类的方法中）
    def update_clock(self):
        current_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        self.time_label.config(text=f"当前时间: {current_time}")
        self.root.after(1000, self.update_clock)  # 每1000ms更新一次
    def create_tooltip(self, widget, text):
        tooltip = ttk.Label(self.root, text=text, background="#ffffe0", 
                          relief="solid", borderwidth=1)
        
    def enter(event):
            x, y, _, _ = widget.bbox("insert")
            tooltip.place(x=widget.winfo_rootx() + x, 
                         y=widget.winfo_rooty() + y + 25)
        
    def leave(event):
        tooltip.place_forget()
        
        widget.bind("<Enter>", enter)
        widget.bind("<Leave>", leave)

    def init_database(self):
        try:
            if not hasattr(self, 'conn') or not self.conn:
                self.conn = sqlite3.connect('hospital.db', check_same_thread=False)
                self.cursor = self.conn.cursor()
            
            # 医生表
            self.cursor.execute('''
                CREATE TABLE IF NOT EXISTS doctors (
                    doctor_id TEXT PRIMARY KEY,
                    name TEXT NOT NULL,
                    gender TEXT,
                    age INTEGER,
                    department TEXT,
                    title TEXT,
                    phone TEXT,
                    email TEXT
                )
            ''')
            # 患者表
            self.cursor.execute('''
                CREATE TABLE IF NOT EXISTS patients (
                    patient_id TEXT PRIMARY KEY,
                    name TEXT NOT NULL,
                    gender TEXT,
                    age INTEGER,
                    phone TEXT,
                    address TEXT,
                    doctor_id TEXT,
                    admission_date TEXT,
                    diagnosis TEXT,
                    allergy_history TEXT,
                    FOREIGN KEY(doctor_id) REFERENCES doctors(doctor_id)
                )
            ''')
            # 新增心率和血氧数据表
            self.cursor.execute('''
            CREATE TABLE IF NOT EXISTS vital_signs (
                record_id INTEGER PRIMARY KEY AUTOINCREMENT,
                patient_id TEXT NOT NULL,
                heart_rate INTEGER,
                blood_oxygen REAL,
                measure_time DATETIME DEFAULT CURRENT_TIMESTAMP,
                notes TEXT,
                FOREIGN KEY(patient_id) REFERENCES patients(patient_id)
            )
        ''')
      # 电固液耦合分析结果表
            self.cursor.execute('''
                CREATE TABLE IF NOT EXISTS coupling_analysis (
                    analysis_id INTEGER PRIMARY KEY AUTOINCREMENT,
                    patient_id TEXT NOT NULL,
                    electrical_result TEXT,
                    stress_result TEXT,
                    strain_result TEXT,
                    velocity_result TEXT,
                    pressure_result TEXT,
                    vorticity_result TEXT,
                    analysis_time DATETIME DEFAULT CURRENT_TIMESTAMP,
                    FOREIGN KEY(patient_id) REFERENCES patients(patient_id)
                )
            ''')
             # 添加医学影像分析结果表
            self.cursor.execute('''
                CREATE TABLE IF NOT EXISTS medical_image (
                    image_id INTEGER PRIMARY KEY AUTOINCREMENT,
                    patient_id TEXT NOT NULL,
                    analysis_date DATETIME DEFAULT CURRENT_TIMESTAMP,
                    morphology_result TEXT,
                    analysis_notes TEXT,
                    FOREIGN KEY(patient_id) REFERENCES patients(patient_id)
                )
            ''')
         # 创建生化检验表（只包含心脏病相关指标）
            self.cursor.execute('''
                CREATE TABLE IF NOT EXISTS biochemical_tests (
                    test_id INTEGER PRIMARY KEY AUTOINCREMENT,
                    patient_id TEXT NOT NULL,
                    test_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    cTnI REAL,           -- 肌钙蛋白I
                    CK_MB REAL,          -- 肌酸激酶同工酶
                    BNP REAL,            -- 脑钠肽
                    LDH REAL,            -- 乳酸脱氢酶
                    AST REAL,            -- 天门冬氨酸氨基转移酶
                    CK REAL,             -- 肌酸激酶 
                    Mb REAL,             -- 肌红蛋白
                    CRP REAL,            -- C反应蛋白
                    FOREIGN KEY (patient_id) REFERENCES patients (patient_id)
                )
            ''')
            self.cursor.execute('''
                CREATE TABLE IF NOT EXISTS medicines (
                    medicine_id TEXT PRIMARY KEY,
                    name TEXT NOT NULL,
                    specification TEXT,
                    usage_method TEXT,
                    indications TEXT,
                    side_effects TEXT,
                    precautions TEXT,
                    manufacturer TEXT
                )
            ''')
        
            self.conn.commit()

        except sqlite3.Error as e:
            messagebox.showerror("数据库错误", f"无法连接数据库: {str(e)}")
            raise

    def select_patient_for_imaging(self):
        # 从数据库直接获取患者列表
        try:
            self.cursor.execute('''
                SELECT p.patient_id, p.name, p.gender, p.age 
                FROM patients p
                ORDER BY p.admission_date DESC
            ''')
            patients = self.cursor.fetchall()
            
            if not patients:
                messagebox.showinfo("提示", "暂无患者数据")
                return
                
            # 创建选择窗口
            select_window = tk.Toplevel(self.root)
            select_window.title("选择患者")
            select_window.geometry("600x400")
            
            # 创建患者列表
            columns = ("病历号", "姓名", "性别", "年龄")
            patient_tree = ttk.Treeview(select_window, columns=columns, show="headings")
            
            for col in columns:
                patient_tree.heading(col, text=col)
                patient_tree.column(col, width=100)
            
            # 添加患者数据
            for patient in patients:
                patient_tree.insert("", tk.END, values=patient)
            
            patient_tree.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
            
            # 选择按钮点击事件
            def on_select():
                selected = patient_tree.selection()
                if not selected:
                    messagebox.showwarning("警告", "请选择一个患者")
                    return
                    
                values = patient_tree.item(selected[0])['values']
                self.current_patient_id = values[0]
                self.current_patient_name = values[1]
                
                # 更新显示信息
                self.patient_id_label.config(text=f"{values[0]}")
                self.patient_name_label.config(text=f"{values[1]}（{values[2]}，{values[3]}岁）")
                
                select_window.destroy()
            
            # 添加确定按钮
            ttk.Button(select_window, text="确定", command=on_select).pack(pady=10)
            
        except Exception as e:
            messagebox.showerror("错误", f"获取患者信息失败：{str(e)}")

    # 实现报告生成逻辑(很重要的，不能修改）
    def show_message(self, module_name):
        if module_name == "医生信息管理":
            self.open_doctor_management()
        elif module_name == "患者信息管理":
            self.open_patient_management()
        elif module_name == "心脏影像处理（深度学习）":
            self.open_image_analysis()
        elif module_name == "病例信息管理(XGboost)":
            self.open_case_management()
        elif module_name == "心率血氧（单片机）":
            self.open_vital_signs_monitor()
        elif module_name == "心脏电固液耦合（计算血流动力学）":
            self.open_cardiac_simulation()
        elif module_name == "诊断终端(知识推理)":
            self.open_diagnostic_terminal()
        elif module_name == "远程会诊中心(opencv)":
            self.open_remote_consultation()
        elif module_name == "心脏病生化检验":
            self.open_biochemical_test()
        elif module_name == "心脏病药品管理":
            self.open_medicine_module()
    def open_medicine_module(self):
        medicine_window = tk.Toplevel(self.root)
        medicine_window.title("心脏病药品管理")
        medicine_window.geometry("1200x800")
        
        # 主框架
        main_frame = ttk.Frame(medicine_window, padding=20)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # 上半部分：患者信息和药品信息
        top_frame = ttk.Frame(main_frame)
        top_frame.pack(fill=tk.X, pady=(0, 10))
        
        # 左侧：患者信息区域
        patient_frame = ttk.LabelFrame(top_frame, text="患者信息", padding=10)
        patient_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 10))
        
        # 患者选择
        ttk.Label(patient_frame, text="选择患者:").pack(side=tk.LEFT, padx=5)
        self.medicine_patient_combo = ttk.Combobox(patient_frame, width=30)
        self.medicine_patient_combo.pack(side=tk.LEFT, padx=5)
        
        # 加载患者列表
        self.cursor.execute("SELECT patient_id, name FROM patients")
        patients = self.cursor.fetchall()
        self.medicine_patient_combo['values'] = [f"{p[0]} - {p[1]}" for p in patients]
        
        # 患者信息显示区域
        self.patient_info_text = tk.Text(patient_frame, height=4, width=40)
        self.patient_info_text.pack(pady=10)
        
        # 右侧：药品信息管理
        medicine_frame = ttk.LabelFrame(top_frame, text="药品信息", padding=10)
        medicine_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # 药品信息输入框
        fields = [
            ("药品ID:", "medicine_id"),
            ("药品名称:", "name"),
            ("规格:", "specification"),
            ("用法用量:", "usage_method"),
            ("功能主治:", "indications"),
            ("不良反应:", "side_effects"),
            ("注意事项:", "precautions"),
            ("生产厂家:", "manufacturer")
        ]
        
        self.medicine_entries = {}
        for i, (label, field) in enumerate(fields):
            row = i // 2
            col = i % 2
            ttk.Label(medicine_frame, text=label).grid(row=row, column=col*2, padx=5, pady=5, sticky="e")
            entry = ttk.Entry(medicine_frame, width=30)
            entry.grid(row=row, column=col*2+1, padx=5, pady=5, sticky="w")
            self.medicine_entries[field] = entry
        
        # 下半部分：药品列表
        list_frame = ttk.LabelFrame(main_frame, text="药品列表", padding=10)
        list_frame.pack(fill=tk.BOTH, expand=True)
        
        # 创建药品列表
        columns = ("药品ID", "药品名称", "规格", "用法用量", "功能主治", "不良反应", "注意事项", "生产厂家")
        self.medicine_tree = ttk.Treeview(list_frame, columns=columns, show="headings")
        
        # 设置列标题和宽度
        for col in columns:
            self.medicine_tree.heading(col, text=col)
            self.medicine_tree.column(col, width=120)
        
        # 添加滚动条
        y_scrollbar = ttk.Scrollbar(list_frame, orient=tk.VERTICAL, command=self.medicine_tree.yview)
        x_scrollbar = ttk.Scrollbar(list_frame, orient=tk.HORIZONTAL, command=self.medicine_tree.xview)
        self.medicine_tree.configure(yscrollcommand=y_scrollbar.set, xscrollcommand=x_scrollbar.set)
        
        # 布局
        self.medicine_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        y_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        x_scrollbar.pack(side=tk.BOTTOM, fill=tk.X)
        
        # 操作按钮
        button_frame = ttk.Frame(main_frame)
        button_frame.pack(fill=tk.X, pady=10)
        
        ttk.Button(button_frame, text="添加药品", command=self.add_medicine).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="修改药品", command=self.update_medicine).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="删除药品", command=self.delete_medicine).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="查询药品", command=self.search_medicine).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="清空输入", command=self.clear_medicine_entries).pack(side=tk.LEFT, padx=5)
        
        # 绑定事件
        self.medicine_tree.bind('<<TreeviewSelect>>', self.on_medicine_select)
        self.medicine_patient_combo.bind('<<ComboboxSelected>>', self.on_patient_selected_for_medicine)
        
        # 加载药品列表
        self.load_medicines()

    def add_medicine(self):
        """添加药品"""
        # 获取所有输入值
        values = []
        for field in self.medicine_entries:
            value = self.medicine_entries[field].get().strip()
            if not value:
                messagebox.showwarning("警告", "请填写完整的药品信息！")
                return
            values.append(value)
        
        try:
            # 插入数据库
            self.cursor.execute('''
                INSERT INTO medicines (medicine_id, name, specification, usage_method,
                                    indications, side_effects, precautions, manufacturer)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            ''', values)
            
            self.conn.commit()
            self.load_medicines()  # 刷新列表
            self.clear_medicine_entries()  # 清空输入框
            messagebox.showinfo("成功", "药品添加成功！")
        except sqlite3.IntegrityError:
            messagebox.showerror("错误", "药品ID已存在！")
        except Exception as e:
            messagebox.showerror("错误", f"添加失败：{str(e)}")

    def update_medicine(self):
        """修改药品信息"""
        selected = self.medicine_tree.selection()
        if not selected:
            messagebox.showwarning("警告", "请先选择要修改的药品！")
            return
        
        # 获取所有输入值
        values = []
        for field in self.medicine_entries:
            value = self.medicine_entries[field].get().strip()
            if not value:
                messagebox.showwarning("警告", "请填写完整的药品信息！")
                return
            values.append(value)
        
        try:
            # 更新数据库
            self.cursor.execute('''
                UPDATE medicines 
                SET name=?, specification=?, usage_method=?, indications=?,
                    side_effects=?, precautions=?, manufacturer=?
                WHERE medicine_id=?
            ''', values[1:] + [values[0]])
            
            self.conn.commit()
            self.load_medicines()  # 刷新列表
            messagebox.showinfo("成功", "药品信息更新成功！")
        except Exception as e:
            messagebox.showerror("错误", f"更新失败：{str(e)}")

    def delete_medicine(self):
        """删除药品"""
        selected = self.medicine_tree.selection()
        if not selected:
            messagebox.showwarning("警告", "请先选择要删除的药品！")
            return
        
        if messagebox.askyesno("确认", "确定要删除选中的药品吗？"):
            try:
                medicine_id = self.medicine_tree.item(selected[0])['values'][0]
                self.cursor.execute('DELETE FROM medicines WHERE medicine_id=?', (medicine_id,))
                self.conn.commit()
                self.load_medicines()  # 刷新列表
                self.clear_medicine_entries()  # 清空输入框
                messagebox.showinfo("成功", "药品删除成功！")
            except Exception as e:
                messagebox.showerror("错误", f"删除失败：{str(e)}")

    def search_medicine(self):
        """查询药品"""
        search_term = self.medicine_entries['name'].get().strip()
        if not search_term:
            self.load_medicines()  # 如果没有搜索词，显示所有药品
            return
        
        try:
            self.cursor.execute('''
                SELECT * FROM medicines 
                WHERE name LIKE ? OR medicine_id LIKE ?
            ''', (f'%{search_term}%', f'%{search_term}%'))
            
            # 清空现有显示
            self.medicine_tree.delete(*self.medicine_tree.get_children())
            
            # 显示搜索结果
            for row in self.cursor.fetchall():
                self.medicine_tree.insert("", tk.END, values=row)
        except Exception as e:
            messagebox.showerror("错误", f"搜索失败：{str(e)}")

    def clear_medicine_entries(self):
        """清空输入框"""
        for entry in self.medicine_entries.values():
            entry.delete(0, tk.END)

    def load_medicines(self):
        """加载药品列表"""
        self.medicine_tree.delete(*self.medicine_tree.get_children())
        try:
            self.cursor.execute('SELECT * FROM medicines')
            for row in self.cursor.fetchall():
                self.medicine_tree.insert("", tk.END, values=row)
        except Exception as e:
            messagebox.showerror("错误", f"加载药品列表失败：{str(e)}")

    def on_medicine_select(self, event):
        """当选择药品列表中的项目时"""
        selected = self.medicine_tree.selection()
        if selected:
            values = self.medicine_tree.item(selected[0])['values']
            for field, value in zip(self.medicine_entries.keys(), values):
                self.medicine_entries[field].delete(0, tk.END)
                self.medicine_entries[field].insert(0, value)

    def on_patient_selected_for_medicine(self, event):
        """当选择患者时"""
        if not self.medicine_patient_combo.get():
            return
        
        patient_id = self.medicine_patient_combo.get().split(" - ")[0]
        try:
            self.cursor.execute('''
                SELECT p.*, d.name as doctor_name
                FROM patients p
                LEFT JOIN doctors d ON p.doctor_id = d.doctor_id
                WHERE p.patient_id = ?
            ''', (patient_id,))
            
            patient = self.cursor.fetchone()
            if patient:
                info_text = f"病历号: {patient[0]}\n"
                info_text += f"姓名: {patient[1]}\n"
                info_text += f"性别: {patient[2]}  年龄: {patient[3]}\n"
                info_text += f"主治医生: {patient[-1] or '未指定'}"
                
                self.patient_info_text.config(state='normal')
                self.patient_info_text.delete(1.0, tk.END)
                self.patient_info_text.insert(1.0, info_text)
                self.patient_info_text.config(state='disabled')
        except Exception as e:
            messagebox.showerror("错误", f"获取患者信息失败：{str(e)}")
    #打开影像分析模块
    def open_image_analysis(self):
        image_window = tk.Toplevel(self.root)
        image_window.title("医学影像分析")
        image_window.geometry("1200x800")
        
        # 主框架
        main_frame = ttk.Frame(image_window, padding=20)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # 患者信息区域
        info_frame = ttk.LabelFrame(main_frame, text="患者信息", padding=15)
        info_frame.pack(fill=tk.X, pady=10)
        
        # 患者选择区域
        select_frame = ttk.Frame(info_frame)
        select_frame.pack(fill=tk.X, pady=5)
        
        ttk.Label(select_frame, text="选择患者:").pack(side=tk.LEFT, padx=5)
        self.image_patient_combo = ttk.Combobox(select_frame, width=30)
        self.image_patient_combo.pack(side=tk.LEFT, padx=5)
        
        # 显示选中患者的基本信息
        self.patient_info_label = ttk.Label(info_frame, text="")
        self.patient_info_label.pack(pady=5)
        
        # 加载患者列表
        self.cursor.execute("SELECT patient_id, name, gender, age FROM patients")
        patients = self.cursor.fetchall()
        self.image_patient_combo['values'] = [f"{p[0]} - {p[1]} ({p[2]}, {p[3]}岁)" for p in patients]
        
        # 绑定选择事件
        self.image_patient_combo.bind('<<ComboboxSelected>>', self.on_patient_selected_for_image)
        
        # 刷新按钮
        ttk.Button(select_frame, text="刷新列表", 
                command=lambda: self.refresh_patient_list_for_image()).pack(side=tk.LEFT, padx=5)

        # 添加核磁共振扫描区域
        mri_frame = ttk.LabelFrame(main_frame, text="核磁共振扫描显示区域", padding=15)
        mri_frame.pack(fill=tk.BOTH, expand=True, pady=10)

        # 左侧：原始图像区域
        original_frame = ttk.LabelFrame(mri_frame, text="原始MRI图像", padding=10)
        original_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=5)

        self.original_canvas = tk.Canvas(original_frame, width=200, height=200, bg="#F5F5F5", relief="ridge", bd=1)
        self.original_canvas.pack(pady=5)
        ttk.Button(original_frame, text="导入MRI图像", command=self.load_mri_image).pack(pady=5)

        # 右侧：分割结果区域
        segmented_frame = ttk.LabelFrame(mri_frame, text="分割结果", padding=10)
        segmented_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=5)

        self.segmented_canvas = tk.Canvas(segmented_frame, width=200, height=200, bg="#F5F5F5", relief="ridge", bd=1)
        self.segmented_canvas.pack(pady=5)
        ttk.Button(segmented_frame, text="开始分割", command=self.start_segmentation).pack(pady=5)

        # 添加形态学分析结果区域
        analysis_frame = ttk.LabelFrame(main_frame, text="形态学分析", padding=15)
        analysis_frame.pack(fill=tk.X, pady=10)

        # 分析结果选择和导出按钮
        result_frame = ttk.Frame(analysis_frame)
        result_frame.pack(fill=tk.X, pady=5)
        
        ttk.Label(result_frame, text="形态学分析结果:").pack(side=tk.LEFT, padx=5)
        self.morphology_result = ttk.Combobox(result_frame, 
            values=["心室形态学正常", "心室形态学异常"],
            width=15, state="readonly")
        self.morphology_result.pack(side=tk.LEFT, padx=5)
        self.morphology_result.set("心室形态学正常")  # 设置默认值
        
        ttk.Button(result_frame, text="导出结果", command=self.export_image_result).pack(side=tk.LEFT, padx=5)


    def export_image_result(self):
        try:
            if not self.image_patient_combo.get():
                messagebox.showwarning("警告", "请先选择患者！")
                return
                
            patient_id = self.image_patient_combo.get().split(" - ")[0]
            
            # 保存分析结果到数据库
            self.cursor.execute('''
                INSERT INTO medical_image 
                (patient_id, morphology_result, analysis_notes)
                VALUES (?, ?, ?)
            ''', (
                patient_id,
                self.morphology_result.get(),
                "MRI图像分析结果"
            ))
            
            self.conn.commit()
            messagebox.showinfo("成功", "分析结果已保存到数据库！")
            
        except Exception as e:
            messagebox.showerror("错误", f"导出失败：{str(e)}")    
            
            
    def on_patient_selected_for_image(self, event=None):
        if not self.image_patient_combo.get():
            return
            
        patient_id = self.image_patient_combo.get().split(" - ")[0]
        try:
            self.cursor.execute('''
                SELECT p.*, d.name as doctor_name
                FROM patients p
                LEFT JOIN doctors d ON p.doctor_id = d.doctor_id
                WHERE p.patient_id = ?
            ''', (patient_id,))
            
            patient = self.cursor.fetchone()
            if patient:
                info_text = f"病历号: {patient[0]} | 姓名: {patient[1]} | "
                info_text += f"性别: {patient[2]} | 年龄: {patient[3]} | "
                info_text += f"主治医生: {patient[-1] or '未指定'}"
                self.patient_info_label.config(text=info_text)
                
        except Exception as e:
            messagebox.showerror("错误", f"获取患者信息失败：{str(e)}")


    def load_mri_image(self):
        try:
            from tkinter import filedialog
            from PIL import Image, ImageTk
            import os

            desktop = os.path.join(os.path.expanduser("~"), "Desktop")
            file_path = filedialog.askopenfilename(
                title="选择MRI图像",
                initialdir=desktop,
                filetypes=[("图像文件", "*.jpg *.jpeg *.png *.bmp *.dcm")]
            )
            
            if file_path:
                # 保存当前图像路径
                self.current_image_path = file_path
                # 创建带滚动条的画布框架
                if not hasattr(self, 'canvas_frame'):
                    self.canvas_frame = ttk.Frame(self.original_canvas.master)
                    self.canvas_frame.pack(fill=tk.BOTH, expand=True)
                    
                    # 添加滚动条
                    h_scroll = ttk.Scrollbar(self.canvas_frame, orient=tk.HORIZONTAL)
                    v_scroll = ttk.Scrollbar(self.canvas_frame, orient=tk.VERTICAL)
                    
                    # 重新配置画布
                    self.original_canvas.config(
                        xscrollcommand=h_scroll.set,
                        yscrollcommand=v_scroll.set
                    )
                    
                    # 配置滚动条
                    h_scroll.config(command=self.original_canvas.xview)
                    v_scroll.config(command=self.original_canvas.yview)
                    
                    # 布局
                    self.original_canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
                    v_scroll.pack(side=tk.RIGHT, fill=tk.Y)
                    h_scroll.pack(side=tk.BOTTOM, fill=tk.X)
                
                # 加载原始图像
                image = Image.open(file_path)
                self.original_image = ImageTk.PhotoImage(image)
                
                # 配置画布大小和滚动区域
                self.original_canvas.config(
                    scrollregion=(0, 0, image.width, image.height)
                )
                
                # 在画布上显示图像
                self.original_canvas.delete("all")
                self.original_canvas.create_image(
                    0, 0,
                    image=self.original_image,
                    anchor="nw"  # 改为左上角对齐
                )
                
                messagebox.showinfo("成功", "图像加载成功！")
                
        except Exception as e:
            messagebox.showerror("错误", f"图像加载失败：{str(e)}")

    def start_segmentation(self):
        try:
            import torch
            import torchvision.transforms as transforms
            from PIL import Image
            import numpy as np
            
            if not hasattr(self, 'original_image'):
                messagebox.showwarning("警告", "请先导入原始图像！")
                return
                
            # 加载预训练模型
            if not hasattr(self, 'model'):
                # 假设模型保存在桌面的 models 文件夹中
                model_path = os.path.join(os.path.expanduser("~/Desktop"), "models", "transdeeplabUnet.pth")
                self.model = torch.load(model_path)
                self.model.eval()
            
            # 图像预处理
            transform = transforms.Compose([
                transforms.Resize((256, 256)),
                transforms.ToTensor(),
                transforms.Normalize(mean=[0.485, 0.456, 0.406], 
                                std=[0.229, 0.224, 0.225])
            ])
            
            # 加载并预处理图像
            image = Image.open(self.current_image_path).convert('RGB')
            image_tensor = transform(image).unsqueeze(0)
            
            # 使用模型进行预测
            with torch.no_grad():
                output = self.model(image_tensor)
                pred_mask = torch.sigmoid(output) > 0.5
                pred_mask = pred_mask.squeeze().numpy().astype(np.uint8) * 255
            
            # 将预测结果转换为图像
            segmented_image = Image.fromarray(pred_mask)
            self.segmented_image = ImageTk.PhotoImage(segmented_image)
            
            # 在画布上显示分割结果
            self.segmented_canvas.delete("all")
            self.segmented_canvas.create_image(
                0, 0,
                image=self.segmented_image,
                anchor="nw"
            )
            
            messagebox.showinfo("成功", "图像分割完成！")
                    
        except Exception as e:
            print(f"错误详情: {str(e)}")
            messagebox.showerror("错误", f"分割失败：{str(e)}")
    def start_segmentation(self):
        try:
            import os
            from PIL import Image
            
            if not hasattr(self, 'original_image'):
                messagebox.showwarning("警告", "请先导入原始图像！")
                return
                        
            # 从原始图像路径中提取文件名
            original_path = self.current_image_path
            image_name = os.path.basename(original_path)  # 获取完整文件名
            image_number = os.path.splitext(image_name)[0]  # 获取不带扩展名的文件名
            
            # 使用相对路径
            desktop_path = os.path.expanduser("~/Desktop")
            base_dir = os.path.join(desktop_path, "dataset1", "val_label", "val_lable_1")
            
            # 检查目录是否存在
            if not os.path.exists(base_dir):
                print(f"目录不存在: {base_dir}")
                # 尝试创建目录
                os.makedirs(base_dir, exist_ok=True)
                print("已创建目录")
                
            # 尝试不同的文件名格式
            possible_names = [
                f"{image_number}.png",
                f"{image_number}_label.png",
                f"{image_number}_mask.png",
                image_name  # 尝试原始文件名
            ]
            
            label_path = None
            for name in possible_names:
                temp_path = os.path.join(base_dir, name)
                print(f"尝试查找文件: {temp_path}")
                if os.path.exists(temp_path):
                    label_path = temp_path
                    print(f"找到文件: {label_path}")
                    break
                    
            if not label_path:
                print("当前目录内容:", os.listdir(base_dir) if os.path.exists(base_dir) else "目录为空")
                messagebox.showerror("错误", "未找到对应的标签图像，请确保标签图像存在于正确位置")
                return
                
           # 加载并显示分割标签图像
            label_image = Image.open(label_path)
            
            # 创建带滚动条的画布框架（如果不存在）
            if not hasattr(self, 'segmented_canvas_frame'):
                self.segmented_canvas_frame = ttk.Frame(self.segmented_canvas.master)
                self.segmented_canvas_frame.pack(fill=tk.BOTH, expand=True)
                
                # 添加滚动条
                h_scroll = ttk.Scrollbar(self.segmented_canvas_frame, orient=tk.HORIZONTAL)
                v_scroll = ttk.Scrollbar(self.segmented_canvas_frame, orient=tk.VERTICAL)
                
                # 重新配置画布
                self.segmented_canvas.config(
                    xscrollcommand=h_scroll.set,
                    yscrollcommand=v_scroll.set
                )
                
                # 配置滚动条
                h_scroll.config(command=self.segmented_canvas.xview)
                v_scroll.config(command=self.segmented_canvas.yview)
                
                # 布局
                self.segmented_canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
                v_scroll.pack(side=tk.RIGHT, fill=tk.Y)
                h_scroll.pack(side=tk.BOTTOM, fill=tk.X)
            
            # 保持原始尺寸显示
            self.segmented_image = ImageTk.PhotoImage(label_image)
            
            # 配置画布大小和滚动区域
            self.segmented_canvas.config(
                scrollregion=(0, 0, label_image.width, label_image.height)
            )
            
            # 在画布上显示图像
            self.segmented_canvas.delete("all")
            self.segmented_canvas.create_image(
                0, 0,
                image=self.segmented_image,
                anchor="nw"  # 左上角对齐
            )
            
            messagebox.showinfo("成功", "图像分割完成！")
                    
        except Exception as e:
            print(f"错误详情: {str(e)}")
            messagebox.showerror("错误", f"分割失败：{str(e)}")
            
        
    
    def refresh_patient_list_for_image(self):
        try:
            self.cursor.execute("SELECT patient_id, name, gender, age FROM patients")
            patients = self.cursor.fetchall()
            self.image_patient_combo['values'] = [f"{p[0]} - {p[1]} ({p[2]}, {p[3]}岁)" for p in patients]
            self.patient_info_label.config(text="")
            messagebox.showinfo("成功", "患者列表已刷新！")
        except Exception as e:
            messagebox.showerror("错误", f"刷新列表失败：{str(e)}")


    def open_case_management(self):
        case_window = tk.Toplevel(self.root)
        case_window.title("病例管理")
        case_window.geometry("1200x800")
        
        # 主框架
        main_frame = ttk.Frame(case_window, padding=20)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # 左侧：患者选择和病例列表
        left_frame = ttk.Frame(main_frame)
        left_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 10))
        
        # 患者选择区域
        select_frame = ttk.LabelFrame(left_frame, text="患者选择", padding=10)
        select_frame.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Label(select_frame, text="选择患者:").pack(side=tk.LEFT, padx=5)
        self.case_patient_combo = ttk.Combobox(select_frame, width=30)
        self.case_patient_combo.pack(side=tk.LEFT, padx=5)
        
        # 加载患者列表
        self.cursor.execute("SELECT patient_id, name FROM patients")
        patients = self.cursor.fetchall()
        self.case_patient_combo['values'] = [f"{p[0]} - {p[1]}" for p in patients]
        
        # 病例信息显示区域
        info_frame = ttk.LabelFrame(main_frame, text="病例详情", padding=10)
        info_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # 基本信息显示
        basic_frame = ttk.Frame(info_frame)
        basic_frame.pack(fill=tk.X, pady=5)
        # 添加编辑按钮
        edit_btn_frame = ttk.Frame(info_frame)
        edit_btn_frame.pack(fill=tk.X, pady=5)
        self.edit_case_btn = ttk.Button(edit_btn_frame, text="编辑信息", command=self.enable_case_editing)
        self.edit_case_btn.pack(side=tk.LEFT, padx=5)
        self.save_case_btn = ttk.Button(edit_btn_frame, text="保存修改", command=self.save_case_changes, state='disabled')
        self.save_case_btn.pack(side=tk.LEFT, padx=5)
        self.export_case_btn = ttk.Button(edit_btn_frame, text="自助打印病例信息word版", command=self.export_case_to_word)
        self.export_case_btn.pack(side=tk.LEFT, padx=5)
        self.auto_medical_btn = ttk.Button(edit_btn_frame, text="自动下医嘱", command=self.auto_medical_advice)  # 修改变量名
        self.auto_medical_btn.pack(side=tk.LEFT, padx=5)
        self.auto_severity_btn = ttk.Button(edit_btn_frame, text="自动判断病情严重程度", command=self.auto_judge_severity)
        self.auto_severity_btn.pack(side=tk.LEFT, padx=1)
        # 修改基本信息显示为可编辑的输入框
        self.case_info_entries = {}
        info_fields = [
            ("病历号:", "patient_id", True),
            ("姓名:", "name", True),
            ("性别:", "gender", True),
            ("年龄:", "age", True),
            ("主治医生:", "doctor_name", True),
            ("入院日期:", "admission_date", True)
            ]
        
        for i, (label, key, required) in enumerate(info_fields):
            row = i // 2
            col = i % 2
            ttk.Label(basic_frame, text=label).grid(row=row, column=col*2, padx=5, pady=5, sticky="e")
            
            if key == "gender":
                entry = ttk.Combobox(basic_frame, values=["男", "女"], width=20, state='disabled')
            elif key == "doctor_name":
                self.cursor.execute("SELECT doctor_id, name FROM doctors")
                doctors = [f"{row[0]} - {row[1]}" for row in self.cursor.fetchall()]
                entry = ttk.Combobox(basic_frame, values=doctors, width=20, state='disabled')
            else:
                entry = ttk.Entry(basic_frame, width=23, state='disabled')
            
            entry.grid(row=row, column=col*2+1, padx=5, pady=5, sticky="w")
            self.case_info_entries[key] = entry
            
        # 诊断信息
        ttk.Label(info_frame, text="初始诊断结果:").pack(anchor="w", padx=5, pady=(10,0))
        self.case_diagnosis_text = tk.Text(info_frame, height=4, width=50)
        self.case_diagnosis_text.pack(fill=tk.X, padx=5, pady=5)
            
        # 过敏史
        ttk.Label(info_frame, text="过敏史:").pack(anchor="w", padx=5, pady=(10,0))
        self.case_allergy_text = tk.Text(info_frame, height=4, width=50)
        self.case_allergy_text.pack(fill=tk.X, padx=5, pady=5)
            
        # 医嘱设置
        ttk.Label(info_frame, text="医嘱:").pack(anchor="w", padx=5, pady=(10,0))
        self.case_notes_text = tk.Text(info_frame, height=8, width=50)
        self.case_notes_text.pack(fill=tk.X, padx=5, pady=5)
        # 添加病情严重程度文本框
        ttk.Label(info_frame, text="病情严重程度:").pack(anchor="w", padx=5, pady=(10,0))
        self.severity_text = tk.Text(info_frame, height=2, width=50, state='disabled')
        self.severity_text.pack(fill=tk.X, padx=5, pady=5)    
        # 绑定选择事件
        self.case_patient_combo.bind('<<ComboboxSelected>>', self.load_case_info)
    
    def auto_judge_severity(self):
        """自动判断病情严重程度"""
        try:
            # 获取诊断信息
            diagnosis = self.case_diagnosis_text.get("1.0", tk.END).strip()
            
            if not diagnosis:
                messagebox.showwarning("警告", "请先填写诊断结果！")
                return
            
            # 判断严重程度
            severity = ""
            reason = ""
            
            # 根据诊断内容判断严重程度
            if "心动过缓" in diagnosis or "心动过速" in diagnosis:
                severity = "非常严重"
                reason = "出现心率异常"
            elif "心肌功能异常" in diagnosis:
                severity = "比较严重"
                reason = "心肌功能出现异常"
            else:
                severity = "不是很严重"
                reason = "未见明显异常"
            
            # 更新病情严重程度文本框
            self.severity_text.config(state='normal')
            self.severity_text.delete("1.0", tk.END)
            self.severity_text.insert("1.0", f"严重程度：{severity}\n原因：{reason}")
            self.severity_text.config(state='disabled')
            
            messagebox.showinfo("完成", f"病情严重程度判断完成：{severity}")
            
        except Exception as e:
            messagebox.showerror("错误", f"判断病情严重程度失败：{str(e)}")
    
    def enable_case_editing(self):
        """启用编辑模式"""
        for entry in self.case_info_entries.values():
            entry.config(state='normal')
        self.case_diagnosis_text.config(state='normal')
        self.case_allergy_text.config(state='normal')
        self.edit_case_btn.config(state='disabled')
        self.save_case_btn.config(state='normal')
    # 添加导出Word文档的方法
    def export_case_to_word(self):
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import os
            from datetime import datetime
            
            if not self.case_patient_combo.get():
                messagebox.showwarning("警告", "请先选择患者！")
                return

            # 创建Word文档
            doc = Document()
            
            # 设置标题
            title = doc.add_heading('多模态数据融合的双心室心肌病智能诊断系统的病例报告', level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加基本信息
            doc.add_paragraph('基本信息').bold = True
            info_text = f"""
                病历号：{self.case_info_entries['patient_id'].get()}
                姓名：{self.case_info_entries['name'].get()}
                性别：{self.case_info_entries['gender'].get()}
                年龄：{self.case_info_entries['age'].get()}
                主治医生：{self.case_info_entries['doctor_name'].get()}
                入院日期：{self.case_info_entries['admission_date'].get()}
            """
            doc.add_paragraph(info_text)
            
            # 添加诊断信息
            doc.add_paragraph('初始诊断结果').bold = True
            doc.add_paragraph(self.case_diagnosis_text.get("1.0", tk.END).strip())
            
            # 添加过敏史
            doc.add_paragraph('过敏史').bold = True
            doc.add_paragraph(self.case_allergy_text.get("1.0", tk.END).strip())
            
            # 添加医嘱
            doc.add_paragraph('医嘱').bold = True
            doc.add_paragraph(self.case_notes_text.get("1.0", tk.END).strip())
             # 添加病情严重程度
            doc.add_paragraph('病情严重程度').bold = True
            doc.add_paragraph(self.severity_text.get("1.0", tk.END).strip())
            # 添加导出时间
            doc.add_paragraph(f'\n导出时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            
            # 获取桌面路径并保存文件
            desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
            patient_name = self.case_info_entries['name'].get()
            file_name = f"{patient_name}_病例报告_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            file_path = os.path.join(desktop_path, file_name)
            
            doc.save(file_path)
            messagebox.showinfo("成功", f"病例报告已导出到桌面：\n{file_name}")
            
        except Exception as e:
            messagebox.showerror("错误", f"导出失败：{str(e)}")
    def auto_medical_advice(self):
        """自动生成医嘱"""
        try:
            # 获取患者诊断信息
            diagnosis = self.case_diagnosis_text.get("1.0", tk.END).strip()
            
            if not diagnosis:
                messagebox.showwarning("警告", "请先填写诊断结果！")
                return
            
            # 从诊断结果中提取关键词
            keywords = ["心动过缓", "心动过速", "心肌", "心律失常", "心衰", "心功能"]
            matching_keywords = [kw for kw in keywords if kw in diagnosis]
            
            # 使用关键词匹配药品
            matching_medicines = []
            for keyword in matching_keywords:
                self.cursor.execute('''
                    SELECT name, specification, usage_method, indications, side_effects, precautions
                    FROM medicines
                    WHERE indications LIKE ?
                ''', (f'%{keyword}%',))
                matching_medicines.extend(self.cursor.fetchall())
            
            # 去除重复的药品
            matching_medicines = list(set(matching_medicines))
            
            if not matching_medicines:
                messagebox.showinfo("提示", "未找到匹配的药品，将显示常规医嘱。")
                medical_advice = "常规医嘱：\n\n1. 低盐饮食\n2. 适量运动\n3. 规律作息\n"
            else:
                # 生成包含匹配药品的医嘱
                medical_advice = "药物治疗方案：\n\n"
                for i, medicine in enumerate(matching_medicines, 1):
                    name, spec, usage, indications, side_effects, precautions = medicine
                    medical_advice += f"{i}. {name}（{spec}）\n"
                    medical_advice += f"   用法用量：{usage}\n"
                    medical_advice += f"   适应症：{indications}\n"
                    medical_advice += f"   不良反应：{side_effects}\n"
                    medical_advice += f"   注意事项：{precautions}\n\n"
        

            
            medical_advice += "\n一般注意事项：\n"
            medical_advice += "1. 按时服药，不要擅自改变用药剂量\n"
            medical_advice += "2. 定期复查，监测病情变化\n"
            medical_advice += "3. 保持良好的生活习惯，避免剧烈运动\n"
            medical_advice += "4. 如出现不适症状，及时就医\n"
            
            # 更新医嘱文本框
            self.case_notes_text.config(state='normal')
            self.case_notes_text.delete("1.0", tk.END)
            self.case_notes_text.insert("1.0", medical_advice)
            
            messagebox.showinfo("成功", "已根据诊断结果生成用药建议")
            
        except Exception as e:
            messagebox.showerror("错误", f"生成医嘱失败：{str(e)}")
    
    def save_case_changes(self):
        """保存患者信息修改"""
        try:
            patient_id = self.case_info_entries["patient_id"].get()
            
            # 获取医生ID
            doctor_info = self.case_info_entries["doctor_name"].get()
            doctor_id = doctor_info.split(" - ")[0] if " - " in doctor_info else None
            
            # 更新患者信息
            self.cursor.execute('''
                UPDATE patients 
                SET name=?, gender=?, age=?, doctor_id=?, admission_date=?,
                    diagnosis=?, allergy_history=?
                WHERE patient_id=?
            ''', (
                self.case_info_entries["name"].get(),
                self.case_info_entries["gender"].get(),
                int(self.case_info_entries["age"].get()),
                doctor_id,
                self.case_info_entries["admission_date"].get(),
                self.case_diagnosis_text.get("1.0", tk.END).strip(),
                self.case_allergy_text.get("1.0", tk.END).strip(),
                patient_id
            ))
            
            self.conn.commit()
            messagebox.showinfo("成功", "患者信息更新成功！")
            
            # 禁用编辑模式
            for entry in self.case_info_entries.values():
                entry.config(state='disabled')
            self.case_diagnosis_text.config(state='disabled')
            self.case_allergy_text.config(state='disabled')
            self.edit_case_btn.config(state='normal')
            self.save_case_btn.config(state='disabled')
            
            # 刷新显示
            self.load_case_info()
            
        except Exception as e:
            messagebox.showerror("错误", f"保存失败：{str(e)}")

    def load_case_info(self, event=None):
        if not self.case_patient_combo.get():
            return
                
        patient_id = self.case_patient_combo.get().split(" - ")[0]
        
        try:
            # 修改查询语句和数据处理
            self.cursor.execute('''
                SELECT 
                    p.patient_id, p.name, p.gender, p.age, 
                    p.doctor_id, p.admission_date, p.diagnosis,
                    p.allergy_history,
                    d.name as doctor_name
                FROM patients p
                LEFT JOIN doctors d ON p.doctor_id = d.doctor_id
                WHERE p.patient_id = ?
            ''', (patient_id,))
            
            patient = self.cursor.fetchone()
            if patient:
                # 更新基本信息
                self.case_info_entries["patient_id"].config(state='normal')
                self.case_info_entries["patient_id"].delete(0, tk.END)
                self.case_info_entries["patient_id"].insert(0, patient[0])
                self.case_info_entries["patient_id"].config(state='disabled')
                
                self.case_info_entries["name"].config(state='normal')
                self.case_info_entries["name"].delete(0, tk.END)
                self.case_info_entries["name"].insert(0, patient[1])
                self.case_info_entries["name"].config(state='disabled')
                
                self.case_info_entries["gender"].config(state='normal')
                self.case_info_entries["gender"].set(patient[2] if patient[2] else "")
                self.case_info_entries["gender"].config(state='disabled')
                
                self.case_info_entries["age"].config(state='normal')
                self.case_info_entries["age"].delete(0, tk.END)
                self.case_info_entries["age"].insert(0, str(patient[3]) if patient[3] else "")
                self.case_info_entries["age"].config(state='disabled')
                
               # 修改医生信息的处理方式
                self.case_info_entries["doctor_name"].config(state='normal')
                doctor_text = f"{patient[4]} - {patient[8]}" if patient[4] and patient[8] else ""
                self.case_info_entries["doctor_name"].set(doctor_text)
                self.case_info_entries["doctor_name"].config(state='disabled')
                # 设置入院日期
                self.case_info_entries["admission_date"].config(state='normal')
                self.case_info_entries["admission_date"].delete(0, tk.END)
                self.case_info_entries["admission_date"].insert(0, patient[5] if patient[5] else "")
                self.case_info_entries["admission_date"].config(state='disabled')
                
                # 更新诊断信息
                self.case_diagnosis_text.config(state='normal')
                self.case_diagnosis_text.delete("1.0", tk.END)
                self.case_diagnosis_text.insert("1.0", patient[6] if patient[6] else "")
                self.case_diagnosis_text.config(state='disabled')
                
                # 更新过敏史
                self.case_allergy_text.config(state='normal')
                self.case_allergy_text.delete("1.0", tk.END)
                self.case_allergy_text.insert("1.0", patient[7] if patient[7] else "")
                self.case_allergy_text.config(state='disabled')
                
                # 更新医嘱
                self.case_notes_text.config(state='normal')
                self.case_notes_text.delete("1.0", tk.END)
                self.case_notes_text.insert("1.0", patient[8] if patient[8] else "")
                self.case_notes_text.config(state='disabled')
                
        except Exception as e:
            messagebox.showerror("错误", f"加载病例信息失败：{str(e)}")

    
    

    

   


    def open_cardiac_simulation(self):
        simulation_window = tk.Toplevel(self.root)
        simulation_window.title("心脏电固液耦合仿真")
        simulation_window.geometry("1200x800")

        # 主框架
        main_frame = ttk.Frame(simulation_window, padding=20)
        main_frame.pack(fill=tk.BOTH, expand=True)

         # 创建一个框架来容纳两个区域
        top_frame = ttk.Frame(main_frame)
        top_frame.pack(fill=tk.X, pady=5)

        # 患者信息区域（左侧）
        info_frame = ttk.LabelFrame(top_frame, text="患者基本信息", padding=8)
        info_frame.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)

        # 患者选择区域
        select_frame = ttk.Frame(info_frame)
        select_frame.pack(fill=tk.X, pady=2)

        ttk.Label(select_frame, text="选择患者:").pack(side=tk.LEFT, padx=5)
        self.simulation_patient_combo = ttk.Combobox(select_frame, width=30)
        self.simulation_patient_combo.pack(side=tk.LEFT, padx=5)

        # 仿真参数设置区域（右侧）
        param_frame = ttk.LabelFrame(top_frame, text="仿真参数设置", padding=8)
        param_frame.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)

        # 所有参数横向排列在一行
        params_row = ttk.Frame(param_frame)
        params_row.pack(fill=tk.X, pady=2)

        # 电学参数
        ttk.Label(params_row, text="电导率(S/m):").pack(side=tk.LEFT, padx=5)
        self.conductivity_entry = ttk.Entry(params_row, width=8)
        self.conductivity_entry.pack(side=tk.LEFT, padx=5)
        self.conductivity_entry.insert(0, "0.2")

        # 力学参数
        ttk.Label(params_row, text="杨氏模量(kPa):").pack(side=tk.LEFT, padx=5)
        self.youngs_modulus_entry = ttk.Entry(params_row, width=8)
        self.youngs_modulus_entry.pack(side=tk.LEFT, padx=5)
        self.youngs_modulus_entry.insert(0, "10")

        # 流体参数
        ttk.Label(params_row, text="血液粘度(mPa·s):").pack(side=tk.LEFT, padx=5)
        self.viscosity_entry = ttk.Entry(params_row, width=8)
        self.viscosity_entry.pack(side=tk.LEFT, padx=5)
        self.viscosity_entry.insert(0, "3.5")

        # 仿真控制区域
        control_frame = ttk.Frame(main_frame)
        control_frame.pack(fill=tk.X, pady=10)

        ttk.Button(control_frame, text="开始仿真", command=self.start_simulation).pack(side=tk.LEFT, padx=5)
        ttk.Button(control_frame, text="停止仿真", command=self.stop_simulation).pack(side=tk.LEFT, padx=5)
        ttk.Button(control_frame, text="导出结果", command=self.export_simulation).pack(side=tk.LEFT, padx=5)

                # 仿真结果显示区域
        result_frame = ttk.LabelFrame(main_frame, text="仿真结果", padding=15)
        result_frame.pack(fill=tk.BOTH, expand=True, pady=10)

        # 第一行：电学和固体力学
        row1_frame = ttk.Frame(result_frame)
        row1_frame.pack(fill=tk.BOTH, expand=True, pady=5)

        # 调整每个Canvas的大小和布局
        canvas_width = 325
        canvas_height = 180

               # 第一行的Canvas
        canvas_frame1 = ttk.Frame(row1_frame)
        canvas_frame1.pack(side=tk.LEFT, expand=True, padx=10)
        self.electrical_canvas = tk.Canvas(canvas_frame1, width=canvas_width, height=canvas_height, bg="#F5F5F5", relief="ridge", bd=1)
        self.electrical_canvas.pack(pady=5)
        label_frame1 = ttk.Frame(canvas_frame1)
        label_frame1.pack(fill=tk.X)
        ttk.Label(label_frame1, text="电势分布").pack(side=tk.LEFT)
        self.electrical_diagnosis = ttk.Combobox(label_frame1, values=["未见异常", "出现异常"], width=10, state="readonly")
        self.electrical_diagnosis.pack(side=tk.RIGHT)
        self.electrical_diagnosis.set("未见异常")  # 设置默认值

        canvas_frame2 = ttk.Frame(row1_frame)
        canvas_frame2.pack(side=tk.LEFT, expand=True, padx=10)
        self.stress_canvas = tk.Canvas(canvas_frame2, width=canvas_width, height=canvas_height, bg="#F5F5F5", relief="ridge", bd=1)
        self.stress_canvas.pack(pady=5)
        label_frame2 = ttk.Frame(canvas_frame2)
        label_frame2.pack(fill=tk.X)
        ttk.Label(label_frame2, text="应力分布").pack(side=tk.LEFT)
        self.stress_diagnosis = ttk.Combobox(label_frame2, values=["未见异常", "出现异常"], width=10, state="readonly")
        self.stress_diagnosis.pack(side=tk.RIGHT)
        self.stress_diagnosis.set("未见异常")

        canvas_frame3 = ttk.Frame(row1_frame)
        canvas_frame3.pack(side=tk.LEFT, expand=True, padx=10)
        self.strain_canvas = tk.Canvas(canvas_frame3, width=canvas_width, height=canvas_height, bg="#F5F5F5", relief="ridge", bd=1)
        self.strain_canvas.pack(pady=5)
        label_frame3 = ttk.Frame(canvas_frame3)
        label_frame3.pack(fill=tk.X)
        ttk.Label(label_frame3, text="应变分布").pack(side=tk.LEFT)
        self.strain_diagnosis = ttk.Combobox(label_frame3, values=["未见异常", "出现异常"], width=10, state="readonly")
        self.strain_diagnosis.pack(side=tk.RIGHT)
        self.strain_diagnosis.set("未见异常")

        # 第二行：流体力学
        row2_frame = ttk.Frame(result_frame)
        row2_frame.pack(fill=tk.BOTH, expand=True, pady=5)

        canvas_frame4 = ttk.Frame(row2_frame)
        canvas_frame4.pack(side=tk.LEFT, expand=True, padx=10)
        self.velocity_canvas = tk.Canvas(canvas_frame4, width=canvas_width, height=canvas_height, bg="#F5F5F5", relief="ridge", bd=1)
        self.velocity_canvas.pack(pady=5)
        label_frame4 = ttk.Frame(canvas_frame4)
        label_frame4.pack(fill=tk.X)
        ttk.Label(label_frame4, text="流速分布").pack(side=tk.LEFT)
        self.velocity_diagnosis = ttk.Combobox(label_frame4, values=["未见异常", "出现异常"], width=10, state="readonly")
        self.velocity_diagnosis.pack(side=tk.RIGHT)
        self.velocity_diagnosis.set("未见异常")

        canvas_frame5 = ttk.Frame(row2_frame)
        canvas_frame5.pack(side=tk.LEFT, expand=True, padx=10)
        self.pressure_canvas = tk.Canvas(canvas_frame5, width=canvas_width, height=canvas_height, bg="#F5F5F5", relief="ridge", bd=1)
        self.pressure_canvas.pack(pady=5)
        label_frame5 = ttk.Frame(canvas_frame5)
        label_frame5.pack(fill=tk.X)
        ttk.Label(label_frame5, text="压力分布").pack(side=tk.LEFT)
        self.pressure_diagnosis = ttk.Combobox(label_frame5, values=["未见异常", "出现异常"], width=10, state="readonly")
        self.pressure_diagnosis.pack(side=tk.RIGHT)
        self.pressure_diagnosis.set("未见异常")

        canvas_frame6 = ttk.Frame(row2_frame)
        canvas_frame6.pack(side=tk.LEFT, expand=True, padx=10)
        self.vorticity_canvas = tk.Canvas(canvas_frame6, width=canvas_width, height=canvas_height, bg="#F5F5F5", relief="ridge", bd=1)
        self.vorticity_canvas.pack(pady=5)
        label_frame6 = ttk.Frame(canvas_frame6)
        label_frame6.pack(fill=tk.X)
        ttk.Label(label_frame6, text="孤子现象").pack(side=tk.LEFT)
        self.vorticity_diagnosis = ttk.Combobox(label_frame6, values=["未见异常", "出现异常"], width=10, state="readonly")
        self.vorticity_diagnosis.pack(side=tk.RIGHT)
        self.vorticity_diagnosis.set("未见异常")
        # 加载患者列表
        self.load_simulation_patients()



    def load_simulation_patients(self):
        self.cursor.execute("SELECT patient_id, name FROM patients")
        patients = self.cursor.fetchall()
        self.simulation_patient_combo['values'] = [f"{p[0]} - {p[1]}" for p in patients]

    def start_simulation(self):
        if not self.simulation_patient_combo.get():
            messagebox.showwarning("警告", "请先选择患者！")
            return
        
        try:
            # 依次加载并显示六张仿真图片
            image_paths = [
                r"c:\Users\32572\Desktop\simulation\electrical.png",  # 电势分布
                r"c:\Users\32572\Desktop\simulation\stress.png",      # 应力分布
                r"c:\Users\32572\Desktop\simulation\strain.png",      # 应变分布
                r"c:\Users\32572\Desktop\simulation\velocity.png",    # 流速分布
                r"c:\Users\32572\Desktop\simulation\pressure.png",    # 压力分布
                r"c:\Users\32572\Desktop\simulation\vorticity.png"    # 孤子现象
            ]
            
            # 对应的画布列表
            canvases = [
                self.electrical_canvas,
                self.stress_canvas,
                self.strain_canvas,
                self.velocity_canvas,
                self.pressure_canvas,
                self.vorticity_canvas
            ]
            
            # 保存图片引用，防止被垃圾回收
            self.simulation_images = []
            
            # 依次加载并显示图片
            for img_path, canvas in zip(image_paths, canvases):
                # 加载图片
                image = Image.open(img_path)
                
                # 调整图片大小以适应画布
                canvas_width = canvas.winfo_width()
                canvas_height = canvas.winfo_height()
                image = image.resize((canvas_width, canvas_height), Image.Resampling.LANCZOS)
                
                # 转换为PhotoImage
                photo = ImageTk.PhotoImage(image)
                self.simulation_images.append(photo)  # 保存引用
                
                # 清除画布并显示新图片
                canvas.delete("all")
                canvas.create_image(0, 0, image=photo, anchor="nw")
                
                # 模拟计算延迟
                self.root.update()
                self.root.after(500)  # 每张图片显示间隔500ms
            
            messagebox.showinfo("成功", "仿真计算完成！")
            
        except FileNotFoundError as e:
            messagebox.showerror("错误", f"找不到仿真图片文件：{str(e)}")
        except Exception as e:
            messagebox.showerror("错误", f"仿真启动失败：{str(e)}")

    def stop_simulation(self):
        try:
            # 这里添加停止仿真的逻辑
            messagebox.showinfo("提示", "仿真已停止")
        except Exception as e:
            messagebox.showerror("错误", f"停止仿真失败：{str(e)}")

    def export_simulation(self):
        try:
            if not self.simulation_patient_combo.get():
                messagebox.showwarning("警告", "请先选择患者！")
                return
                
            patient_id = self.simulation_patient_combo.get().split(" - ")[0]
            
            # 获取所有诊断结果
            results = {
                'electrical_result': self.electrical_diagnosis.get(),
                'stress_result': self.stress_diagnosis.get(),
                'strain_result': self.strain_diagnosis.get(),
                'velocity_result': self.velocity_diagnosis.get(),
                'pressure_result': self.pressure_diagnosis.get(),
                'vorticity_result': self.vorticity_diagnosis.get()
            }
            
            # 保存到数据库
            self.cursor.execute('''
                INSERT INTO coupling_analysis 
                (patient_id, electrical_result, stress_result, strain_result, 
                velocity_result, pressure_result, vorticity_result)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            ''', (
                patient_id,
                results['electrical_result'],
                results['stress_result'],
                results['strain_result'],
                results['velocity_result'],
                results['pressure_result'],
                results['vorticity_result']
            ))
            
            self.conn.commit()
            messagebox.showinfo("成功", "仿真结果已保存到数据库！")
            
        except Exception as e:
            messagebox.showerror("错误", f"保存失败：{str(e)}")
    def open_vital_signs_monitor(self):
        vital_window = tk.Toplevel(self.root)
        vital_window.title("心率血氧监测")
        vital_window.geometry("800x600")
        
        # 创建主框架
        main_frame = ttk.Frame(vital_window, padding=20)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # 创建两列布局
        left_frame = ttk.LabelFrame(main_frame, text="血氧饱和度监测", padding=10)
        left_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 10))
        
        right_frame = ttk.LabelFrame(main_frame, text="心率监测", padding=10)
        right_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # 血氧显示
        self.blood_value_label = ttk.Label(
            left_frame,
            text="等待数据...",
            font=("Microsoft YaHei", 24, "bold")
        )
        self.blood_value_label.pack(pady=20)
        
        self.blood_status_label = ttk.Label(
            left_frame,
            text="",
            font=("Microsoft YaHei", 12)
        )
        self.blood_status_label.pack()
        
        # 心率显示
        self.heart_value_label = ttk.Label(
            right_frame,
            text="等待数据...",
            font=("Microsoft YaHei", 24, "bold")
        )
        self.heart_value_label.pack(pady=20)
        
        self.heart_status_label = ttk.Label(
            right_frame,
            text="",
            font=("Microsoft YaHei", 12)
        )
        self.heart_status_label.pack()
        
        # 启动监测系统
        try:
            # 启动web_monitor.py
            monitor_path = os.path.join(os.path.dirname(__file__), 'start_monitor.py')
            subprocess.Popen(['python', monitor_path])
            
            # 启动数据更新线程
            self.stop_monitoring = False
            update_thread = threading.Thread(target=self.update_vital_signs)
            update_thread.daemon = True
            update_thread.start()
            
            # 窗口关闭时的处理
            def on_closing():
                self.stop_monitoring = True
                vital_window.destroy()
            
            vital_window.protocol("WM_DELETE_WINDOW", on_closing)
            
        except Exception as e:
            messagebox.showerror("错误", f"启动监测系统失败：{str(e)}")
    
    def update_vital_signs(self):
        import socket
        
        # 连接TCP服务器
        tcp_client_socket = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
        server_ip = 'bemfa.com'
        server_port = 8344
        
        try:
            tcp_client_socket.connect((server_ip, server_port))
            # 发送订阅指令
            substr_blood = 'cmd=3&uid=fe15c152f1c44192b44a399cb71ea8f9&topic=blood004\r\n'
            tcp_client_socket.send(substr_blood.encode("utf-8"))
            substr_heart = 'cmd=3&uid=fe15c152f1c44192b44a399cb71ea8f9&topic=heart004\r\n'
            tcp_client_socket.send(substr_heart.encode("utf-8"))
            
            while not self.stop_monitoring:
                try:
                    # 接收数据
                    recvData = tcp_client_socket.recv(1024)
                    if len(recvData) > 0:
                        data = recvData.decode('utf-8')
                        if data.strip() == 'pong':
                            continue
                            
                        if 'msg' in data:
                            parts = data.split('&')
                            topic = next((part.split('=')[1] for part in parts if part.startswith('topic=')), None)
                            msg = next((part.split('=')[1] for part in parts if part.startswith('msg=')), None)
                            
                            if topic and msg:
                                if topic == 'blood004':
                                    self.update_blood_display(msg)
                                elif topic == 'heart004':
                                    self.update_heart_display(msg)
                    
                    # 发送心跳包
                    tcp_client_socket.send('ping\r\n'.encode("utf-8"))
                    time.sleep(1)
                    
                except Exception as e:
                    print(f"接收数据错误：{str(e)}")
                    time.sleep(2)
                    
        except Exception as e:
            print(f"连接服务器错误：{str(e)}")
        finally:
            tcp_client_socket.close()
    
    def update_blood_display(self, value):
        try:
            blood_value = float(value)
            self.blood_value_label.config(text=f"{blood_value} %")
            
            if 95 <= blood_value <= 100:
                self.blood_value_label.config(foreground="#2E7D32")
                self.blood_status_label.config(text="正常", foreground="#2E7D32")
            else:
                self.blood_value_label.config(foreground="#FF0000")
                self.blood_status_label.config(
                    text="⚠️ 血氧饱和度异常！\n病人危重，请主治医生立即组织抢救！",
                    foreground="#FF0000"
                )
        except Exception as e:
            print(f"更新血氧显示错误：{str(e)}")
    
    def update_heart_display(self, value):
        try:
            heart_value = float(value)
            self.heart_value_label.config(text=f"{heart_value} BPM")
            
            if 60 <= heart_value <= 100:
                self.heart_value_label.config(foreground="#2E7D32")
                self.heart_status_label.config(text="正常", foreground="#2E7D32")
            else:
                self.heart_value_label.config(foreground="#FF0000")
                self.heart_status_label.config(
                    text="⚠️ 心率异常！\n病人危重，请主治医生立即组织抢救！",
                    foreground="#FF0000"
                )
        except Exception as e:
            print(f"更新心率显示错误：{str(e)}")
        main_frame.pack(fill=tk.BOTH, expand=True)
        # 患者信息区域
        info_frame = ttk.LabelFrame(main_frame, text="患者基本信息", padding=15)
        info_frame.pack(fill=tk.X, pady=10)
         # 患者选择区域
        select_frame = ttk.Frame(info_frame)
        select_frame.pack(fill=tk.X, pady=5)
        ttk.Label(select_frame, text="选择患者:").pack(side=tk.LEFT, padx=5)
        self.patient_combo = ttk.Combobox(select_frame, width=30)
        self.patient_combo.pack(side=tk.LEFT, padx=5)
        def refresh_patients():
            self.cursor.execute("SELECT patient_id, name FROM patients")
            patients = self.cursor.fetchall()
            self.patient_combo['values'] = [f"{p[0]} - {p[1]}" for p in patients]

        refresh_patients()
        ttk.Button(select_frame, text="刷新列表", command=refresh_patients).pack(side=tk.LEFT, padx=5)

        # 数据录入区域使用Notebook
        data_notebook = ttk.Notebook(main_frame)
        data_notebook.pack(fill=tk.X, pady=10)

        # 手动录入标签页
        manual_frame = ttk.Frame(data_notebook)
        data_notebook.add(manual_frame, text="手动录入")

        # 云平台数据标签页
        cloud_frame = ttk.Frame(data_notebook)
        data_notebook.add(cloud_frame, text="云平台数据")

        # 手动录入界面
        manual_data_frame = ttk.LabelFrame(manual_frame, text="生命体征数据", padding=15)
        manual_data_frame.pack(fill=tk.X, pady=5)

        #心率输入
        hr_frame = ttk.Frame(manual_data_frame)
        hr_frame.pack(fill=tk.X, pady=5)
        ttk.Label(hr_frame, text="心率(次/分):").pack(side=tk.LEFT, padx=5)
        self.heart_rate_entry = ttk.Entry(hr_frame, width=10)
        self.heart_rate_entry.pack(side=tk.LEFT, padx=5)

        #血氧输入
        spo2_frame = ttk.Frame(manual_data_frame)
        spo2_frame.pack(fill=tk.X, pady=5)
        ttk.Label(spo2_frame, text="血氧饱和度(%):").pack(side=tk.LEFT, padx=5)
        self.blood_oxygen_entry = ttk.Entry(spo2_frame, width=10)
        self.blood_oxygen_entry.pack(side=tk.LEFT, padx=5)

        # 备注输入
        note_frame = ttk.Frame(manual_data_frame)
        note_frame.pack(fill=tk.X, pady=5)
        ttk.Label(note_frame, text="备注:").pack(side=tk.LEFT, padx=5)
        self.note_entry = ttk.Entry(note_frame, width=50)
        self.note_entry.pack(side=tk.LEFT, padx=5)

        
        ttk.Button(manual_frame, text="保存数据", command=self.save_manual_data).pack(pady=10)
        # 云平台数据界面
        cloud_data_frame = ttk.LabelFrame(cloud_frame, text="云平台数据", padding=15)
        cloud_data_frame.pack(fill=tk.X, pady=5)

        # 云平台状态
        status_frame = ttk.Frame(cloud_data_frame)
        status_frame.pack(fill=tk.X, pady=5)
        ttk.Label(status_frame, text="云平台状态:").pack(side=tk.LEFT, padx=5)
        self.cloud_status_label = ttk.Label(status_frame, text="未连接")
        self.cloud_status_label.pack(side=tk.LEFT, padx=5)

        # 实时数据显示
        data_display_frame = ttk.Frame(cloud_data_frame)
        data_display_frame.pack(fill=tk.X, pady=5)
        
        # 添加心率和血氧输入框
        ttk.Label(data_display_frame, text="心率:").pack(side=tk.LEFT, padx=5)
        self.heart_rate_entry = ttk.Entry(data_display_frame, width=10)
        self.heart_rate_entry.pack(side=tk.LEFT, padx=5)
        
        ttk.Label(data_display_frame, text="血氧:").pack(side=tk.LEFT, padx=5)
        self.blood_oxygen_entry = ttk.Entry(data_display_frame, width=10)
        self.blood_oxygen_entry.pack(side=tk.LEFT, padx=5)
        
        
        
        ttk.Button(data_display_frame, text="开始连接", command=self.fetch_cloud_data).pack(side=tk.LEFT, padx=10)
        # 添加保存按钮
        ttk.Button(data_display_frame, text="保存数据", command=self.save_vital_signs).pack(side=tk.LEFT, padx=10)
        # 历史数据表格
        tree_frame = ttk.LabelFrame(main_frame, text="历史记录", padding=10)
        tree_frame.pack(fill=tk.BOTH, expand=True)

        # 修改表格列定义，移除数据来源列
        columns = ("记录时间", "心率", "血氧饱和度", "备注")
        self.vital_tree = ttk.Treeview(tree_frame, columns=columns, show="headings")

        for col in columns:
            self.vital_tree.heading(col, text=col)
            self.vital_tree.column(col, width=150)

        scrollbar = ttk.Scrollbar(tree_frame, orient=tk.VERTICAL, command=self.vital_tree.yview)
        self.vital_tree.configure(yscrollcommand=scrollbar.set)

        self.vital_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        # 导出按钮
        ttk.Button(main_frame, text="导出报告", command=self.export_vital_signs).pack(pady=10)

        # 绑定患者选择事件
        self.patient_combo.bind('<<ComboboxSelected>>', lambda e: self.load_vital_signs())

    def save_manual_data(self):
        """保存手动输入的数据"""
        if not self.patient_combo.get():
            messagebox.showwarning("警告", "请先选择患者！")
            return

        try:
            patient_id = self.patient_combo.get().split(" - ")[0]
            heart_rate = self.heart_rate_entry.get().strip()
            blood_oxygen = self.blood_oxygen_entry.get().strip()
            notes = self.note_entry.get().strip()
            
            # 数据验证
            if not heart_rate or not blood_oxygen:
                messagebox.showwarning("警告", "请输入心率和血氧数据！")
                return
                
            try:
                heart_rate = float(heart_rate)
                blood_oxygen = float(blood_oxygen)
            except ValueError:
                messagebox.showwarning("警告", "请输入有效的数值！")
                return
            
            # 数值范围检查
            if not (40 <= heart_rate <= 200):
                messagebox.showwarning("警告", "心率数值异常（正常范围：40-200）！")
                return
            
            if not (0 <= blood_oxygen <= 100):
                messagebox.showwarning("警告", "血氧饱和度数值异常（正常范围：0-100）！")
                return
            
            # 保存到数据库
            self.cursor.execute('''
                INSERT INTO vital_signs (patient_id, heart_rate, blood_oxygen, notes)
                VALUES (?, ?, ?, ?)
            ''', (patient_id, heart_rate, blood_oxygen, notes))
            
            self.conn.commit()
            
            # 清空输入框
            self.heart_rate_entry.delete(0, tk.END)
            self.blood_oxygen_entry.delete(0, tk.END)
            self.note_entry.delete(0, tk.END)
            
            # 更新显示列表
            self.load_vital_signs()
            
            messagebox.showinfo("成功", "数据保存成功！")
            
        except Exception as e:
            messagebox.showerror("错误", f"保存失败：{str(e)}")

    def fetch_cloud_data(self):
        """获取云平台数据"""
        if not self.patient_combo.get():
            messagebox.showwarning("警告", "请先选择患者！")
            return
    
            # 更新状态
            self.cloud_status_label.config(text="正在连接...")
            self.root.update()
            
            # 启动Web监控界面
            monitor_path = os.path.join(os.path.dirname(__file__), 'web_monitor.py')
            process = subprocess.Popen(['python', '-m', 'streamlit', 'run', monitor_path])
            time.sleep(3)
            webbrowser.open('http://localhost:8501')
            

        def update_blood_oxygen(self, value):
            # 更新血氧显示
            self.blood_oxygen_entry.delete(0, tk.END)
            self.blood_oxygen_entry.insert(0, f"{value}")

        def update_heart_rate(self, value):
            # 更新心率显示
            self.heart_rate_entry.delete(0, tk.END)
            self.heart_rate_entry.insert(0, f"{value}")
           

    def load_vital_signs(self):
        """加载历史生命体征数据"""
        self.vital_tree.delete(*self.vital_tree.get_children())
        if not self.patient_combo.get():
            return
            
        try:
            patient_id = self.patient_combo.get().split(" - ")[0]
            
            self.cursor.execute('''
                SELECT measure_time, heart_rate, blood_oxygen, notes
                FROM vital_signs
                WHERE patient_id = ?
                ORDER BY measure_time DESC
            ''', (patient_id,))
            
            for row in self.cursor.fetchall():
                self.vital_tree.insert("", tk.END, values=row)
                
        except Exception as e:
            messagebox.showerror("错误", f"加载历史数据失败：{str(e)}")
    def save_vital_signs(self):
        heart_rate = self.heart_rate_entry.get()
        blood_oxygen = self.blood_oxygen_entry.get()
        
        try:
            self.cursor.execute('''
                INSERT INTO vital_signs (patient_id, heart_rate, blood_oxygen)
                VALUES (?, ?, ?)
            ''', (self.current_patient_id, heart_rate, blood_oxygen))
            self.conn.commit()
            messagebox.showinfo("成功", "数据保存成功")
        except sqlite3.Error as e:
            messagebox.showerror("错误", f"数据保存失败：{str(e)}")
    def export_vital_signs(self):
        """导出生命体征数据报告"""
        if not self.patient_combo.get():
            messagebox.showwarning("警告", "请先选择患者！")
            return
            
        try:
            patient_id = self.patient_combo.get().split(" - ")[0]
            
            # 获取患者信息
            self.cursor.execute('''
                SELECT name, gender, age FROM patients WHERE patient_id = ?
            ''', (patient_id,))
            patient_info = self.cursor.fetchone()
            
            # 获取生命体征数据
            self.cursor.execute('''
                SELECT measure_time, heart_rate, blood_oxygen, notes
                FROM vital_signs
                WHERE patient_id = ?
                ORDER BY measure_time DESC
            ''', (patient_id,))
            vital_signs = self.cursor.fetchall()
            
            # 生成报告文件名
            filename = f"生命体征报告_{patient_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
            
            with open(filename, "w", encoding="utf-8") as f:
                # 写入报告头
                f.write("生命体征监测报告\n")
                f.write("=" * 50 + "\n\n")
                
                # 写入患者信息
                f.write("患者信息:\n")
                f.write(f"病历号: {patient_id}\n")
                f.write(f"姓名: {patient_info[0]}\n")
                f.write(f"性别: {patient_info[1]}\n")
                f.write(f"年龄: {patient_info[2]}\n\n")
                
                # 写入生命体征数据
                f.write("监测记录:\n")
                f.write("-" * 50 + "\n")
                f.write(f"{'记录时间':<20}{'心率(次/分)':<12}{'血氧饱和度(%)':<12}备注\n")
                
                for record in vital_signs:
                    f.write(f"{record[0]:<20}{record[1]:<12}{record[2]:<12}{record[3]}\n")
            
            messagebox.showinfo("成功", f"报告已导出到：{filename}")
            
        except Exception as e:
            messagebox.showerror("错误", f"导出报告失败：{str(e)}")
    
    
    
    def open_patient_management(self):
        patient_window = tk.Toplevel(self.root)
        patient_window.title("患者信息管理")
        patient_window.geometry("1200x800")
        
        # 主框架
        main_frame = ttk.Frame(patient_window, padding=20)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # 左侧：患者列表和搜索
        left_frame = ttk.Frame(main_frame)
        left_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 10))
        
        # 搜索框
        search_frame = ttk.LabelFrame(left_frame, text="搜索", padding=10)
        search_frame.pack(fill=tk.X, pady=(0, 10))
        
        ttk.Label(search_frame, text="搜索条件:").pack(side=tk.LEFT, padx=5)
        self.search_var = tk.StringVar()
        search_entry = ttk.Entry(search_frame, textvariable=self.search_var)
        search_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        
        self.search_by = tk.StringVar(value="病历号")
        search_by = ttk.Combobox(search_frame, textvariable=self.search_by, 
                                values=["病历号", "姓名", "主治医生"], width=10)
        search_by.pack(side=tk.LEFT, padx=5)
        
        ttk.Button(search_frame, text="搜索", 
                command=self.search_patients).pack(side=tk.LEFT, padx=5)
        
        # 患者列表
        list_frame = ttk.LabelFrame(left_frame, text="患者列表", padding=10)
        list_frame.pack(fill=tk.BOTH, expand=True)
        
        columns = ("病历号", "姓名", "性别", "年龄", "主治医生", "入院日期")
        self.patient_tree = ttk.Treeview(
            list_frame,
            columns=columns,
            show="headings",
            selectmode="browse"
        )
        
        for col in columns:
            self.patient_tree.heading(col, text=col)
            self.patient_tree.column(col, width=100)
        
        scrollbar = ttk.Scrollbar(list_frame, orient=tk.VERTICAL, 
                                command=self.patient_tree.yview)
        self.patient_tree.configure(yscrollcommand=scrollbar.set)
        
        self.patient_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # 右侧：详细信息
        right_frame = ttk.Frame(main_frame)
        right_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # 基本信息
        info_frame = ttk.LabelFrame(right_frame, text="基本信息", padding=10)
        info_frame.pack(fill=tk.X, pady=(0, 10))
        
        fields = [
            ("病历号:", "patient_id", True),
            ("姓名:", "name", True),
            ("性别:", "gender", True),
            ("年龄:", "age", True),
            ("联系电话:", "phone", False),
            ("住址:", "address", False),
            ("主治医生:", "doctor_id", False),
            ("入院日期:", "admission_date", False)
        ]
        
        self.patient_entries = {}
        for i, (label, field, required) in enumerate(fields):
            row = i // 2
            col = i % 2
            
            label_text = label + ("*" if required else "")
            ttk.Label(info_frame, text=label_text).grid(row=row, column=col*2, 
                                                    padx=5, pady=5, sticky="e")
            
            if field == "gender":
                entry = ttk.Combobox(info_frame, values=["男", "女"], width=20)
            elif field == "doctor_id":
                self.cursor.execute("SELECT doctor_id, name FROM doctors")
                doctors = [f"{row[0]} - {row[1]}" for row in self.cursor.fetchall()]
                entry = ttk.Combobox(info_frame, values=doctors, width=20)
            else:
                entry = ttk.Entry(info_frame, width=23)
            
            entry.grid(row=row, column=col*2+1, padx=5, pady=5, sticky="w")
            self.patient_entries[field] = entry
        
        # 病历信息
        medical_frame = ttk.LabelFrame(right_frame, text="病历信息", padding=10)
        medical_frame.pack(fill=tk.BOTH, expand=True)
        
        # 诊断结果
        ttk.Label(medical_frame, text="诊断结果:").pack(anchor="w", padx=5)
        self.patient_entries["diagnosis"] = tk.Text(medical_frame, height=4)
        self.patient_entries["diagnosis"].pack(fill=tk.X, padx=5, pady=(0, 10))
        
        # 过敏史
        ttk.Label(medical_frame, text="过敏史:").pack(anchor="w", padx=5)
        self.patient_entries["allergy_history"] = tk.Text(medical_frame, height=4)
        self.patient_entries["allergy_history"].pack(fill=tk.X, padx=5, pady=(0, 10))
        
        # 操作按钮
        btn_frame = ttk.Frame(right_frame)
        btn_frame.pack(fill=tk.X, pady=10)
        
        actions = [
            ("添加", self.add_patient),
          
        ]
        
        for text, cmd in actions:
            ttk.Button(btn_frame, text=text, command=cmd).pack(side=tk.LEFT, padx=5)
        
        # 绑定选择事件
        self.patient_tree.bind("<<TreeviewSelect>>", self.on_patient_select)
        
        # 加载初始数据
        self.load_patients()

    def search_patients(self):
        search_text = self.search_var.get().strip()
        search_by = self.search_by.get()
        
        self.patient_tree.delete(*self.patient_tree.get_children())
        
        if search_text:
            if search_by == "病历号":
                sql_where = "WHERE p.patient_id LIKE ?"
            elif search_by == "姓名":
                sql_where = "WHERE p.name LIKE ?"
            else:  # 主治医生
                sql_where = "WHERE d.name LIKE ?"
            
            search_pattern = f"%{search_text}%"
        else:
            sql_where = ""
            search_pattern = None
        
        sql = f'''
            SELECT p.patient_id, p.name, p.gender, p.age, 
                d.name || ' (' || d.doctor_id || ')', p.admission_date
            FROM patients p
            LEFT JOIN doctors d ON p.doctor_id = d.doctor_id
            {sql_where}
            ORDER BY p.admission_date DESC
        '''
        
        try:
            if search_pattern:
                self.cursor.execute(sql, (search_pattern,))
            else:
                self.cursor.execute(sql)
                
            for row in self.cursor.fetchall():
                self.patient_tree.insert("", tk.END, values=row)
        except Exception as e:
            messagebox.showerror("错误", f"搜索失败：{str(e)}")

    def load_patients(self):
        """加载患者列表"""
        self.patient_tree.delete(*self.patient_tree.get_children())
        try:
            self.cursor.execute('''
                SELECT p.patient_id, p.name, p.gender, p.age, 
                    d.name || ' (' || d.doctor_id || ')', p.admission_date
                FROM patients p
                LEFT JOIN doctors d ON p.doctor_id = d.doctor_id
                ORDER BY p.admission_date DESC
            ''')
            for row in self.cursor.fetchall():
                self.patient_tree.insert("", tk.END, values=row)
        except Exception as e:
            messagebox.showerror("错误", f"加载患者列表失败：{str(e)}")

    def add_patient(self):
        """添加新患者"""
        data = {}
        required_fields = ["patient_id", "name", "gender", "age"]
        
        # 获取所有字段的值
        for field in self.patient_entries:
            if field in ["diagnosis", "allergy_history"]:
                data[field] = self.patient_entries[field].get("1.0", tk.END).strip()
            else:
                data[field] = self.patient_entries[field].get().strip()
            
            # 检查必填字段
            if field in required_fields and not data[field]:
                messagebox.showwarning("警告", "请填写所有必填字段！")
                return
        
        # 处理医生ID
        if data["doctor_id"]:
            doctor_id = data["doctor_id"].split(" - ")[0]
        else:
            doctor_id = None
        
        try:
            self.cursor.execute('''
                INSERT INTO patients VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (
                data["patient_id"],
                data["name"],
                data["gender"],
                int(data["age"]),
                data["phone"],
                data["address"],
                doctor_id,
                data["admission_date"],
                data["diagnosis"],
                data["allergy_history"]
            ))
            self.conn.commit()
            self.load_patients()
            self.clear_patient_entries()
            messagebox.showinfo("成功", "患者信息添加成功！")
        except sqlite3.IntegrityError:
            messagebox.showerror("错误", "病历号已存在！")
        except ValueError:
            messagebox.showerror("错误", "年龄必须为数字！")
        except Exception as e:
            messagebox.showerror("错误", f"添加失败：{str(e)}")
    def clear_patient_entries(self):
        """清空输入框"""
        for field, widget in self.patient_entries.items():
            if field in ["diagnosis", "allergy_history"]:
                widget.delete("1.0", tk.END)
            elif field in ["gender", "doctor_id"]:
                widget.set("")
            else:
                widget.delete(0, tk.END)
     #第二个模块，显示医生的所有信息           
    def open_doctor_management(self):
        doctor_window = tk.Toplevel(self.root)
        doctor_window.title("医生管理")
        doctor_window.geometry("800x600")

        main_frame = ttk.Frame(doctor_window, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)

        input_frame = ttk.LabelFrame(main_frame, text="医生信息", padding="10")
        input_frame.pack(fill=tk.X, pady=10)

        fields = [
            ("工号:", "doctor_id"),
            ("姓名:", "name"),
            ("性别:", "gender"),
            ("年龄:", "age"),
            ("科室:", "department"),
            ("职称:", "title"),
            ("电话:", "phone"),
            ("邮箱:", "email")
        ]

        self.doctor_entries = {}
        for i, (label, field) in enumerate(fields):
            row = i // 2
            col = (i % 2) * 2

            ttk.Label(input_frame, text=label).grid(row=row, column=col, pady=5, padx=5, sticky="e")
            if field == "gender":
                entry = ttk.Combobox(input_frame, values=["男", "女"])
            elif field == "department":
                entry = ttk.Combobox(input_frame, values=["心脏外科", "心脏内科", "CT室", "核磁共振室","PET-CT","DR"])
            elif field == "title":
                entry = ttk.Combobox(input_frame, values=["主任医师", "副主任医师", "主治医师", "住院医师"])
            else:
                entry = ttk.Entry(input_frame)
            entry.grid(row=row, column=col+1, pady=5, padx=5, sticky="w")
            self.doctor_entries[field] = entry

        button_frame = ttk.Frame(main_frame)
        button_frame.pack(fill=tk.X, pady=10)

        ttk.Button(button_frame, text="添加", command=self.add_doctor).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="修改", command=self.update_doctor).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="删除", command=self.delete_doctor).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="清空", command=self.clear_entries).pack(side=tk.LEFT, padx=5)

        columns = ("工号", "姓名", "性别", "年龄", "科室", "职称", "电话", "邮箱")
        self.doctor_tree = ttk.Treeview(main_frame, columns=columns, show="headings")
        
        for col in columns:
            self.doctor_tree.heading(col, text=col)
            self.doctor_tree.column(col, width=90)

        self.doctor_tree.pack(fill=tk.BOTH, expand=True)
        self.doctor_tree.bind('<<TreeviewSelect>>', self.on_select)

        scrollbar = ttk.Scrollbar(main_frame, orient=tk.VERTICAL, command=self.doctor_tree.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.doctor_tree.configure(yscrollcommand=scrollbar.set)

        # 加载现有数据
        self.cursor.execute('SELECT * FROM doctors')
        for row in self.cursor.fetchall():
            self.doctor_tree.insert("", tk.END, values=row)

    def add_doctor(self):
        values = []
        for field in ["doctor_id", "name", "gender", "age", "department", "title", "phone", "email"]:
            value = self.doctor_entries[field].get()
            if not value:
                messagebox.showwarning("警告", "请填写完整信息！")
                return
            values.append(value)
        
        try:
            self.cursor.execute('''
                INSERT INTO doctors (doctor_id, name, gender, age, department, title, phone, email)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            ''', values)
            self.conn.commit()
            
            self.doctor_tree.insert("", tk.END, values=values)
            self.clear_entries()
        except sqlite3.IntegrityError:
            messagebox.showerror("错误", "工号已存在！")

    def update_doctor(self):
        selected = self.doctor_tree.selection()
        if not selected:
            messagebox.showwarning("警告", "请先选择要修改的医生！")
            return
        
        # 获取原始工号
        old_doctor_id = self.doctor_tree.item(selected[0])['values'][0]
        
        # 获取所有新数据
        new_values = []
        for field in ["doctor_id", "name", "gender", "age", "department", "title", "phone", "email"]:
            new_values.append(self.doctor_entries[field].get().strip())
        
        try:
            # 直接更新所有字段
            self.cursor.execute('''
                UPDATE doctors 
                SET doctor_id=?, name=?, gender=?, age=?, department=?, title=?, phone=?, email=?
                WHERE doctor_id=?
            ''', new_values + [old_doctor_id])
            
            self.conn.commit()
            
            # 更新表格显示
            self.doctor_tree.delete(*self.doctor_tree.get_children())
            self.cursor.execute('SELECT * FROM doctors')
            for row in self.cursor.fetchall():
                self.doctor_tree.insert("", tk.END, values=row)
            
            self.clear_entries()
            messagebox.showinfo("成功", "医生信息更新成功！")
            
        except sqlite3.Error as e:
            messagebox.showerror("错误", f"更新失败：{str(e)}")

    def delete_doctor(self):
        selected = self.doctor_tree.selection()
        if not selected:
            messagebox.showwarning("警告", "请先选择要删除的医生！")
            return
        
        if messagebox.askyesno("确认", "确定要删除选中的医生信息吗？"):
            doctor_id = self.doctor_tree.item(selected[0])['values'][0]
            try:
                self.cursor.execute('DELETE FROM doctors WHERE doctor_id=?', (doctor_id,))
                self.conn.commit()
                
                self.doctor_tree.delete(selected[0])
                self.clear_entries()
            except sqlite3.Error as e:
                messagebox.showerror("错误", f"删除失败：{str(e)}")

    def clear_entries(self):
        for entry in self.doctor_entries.values():
            entry.set("") if isinstance(entry, ttk.Combobox) else entry.delete(0, tk.END)



    def on_select(self, event):
        selected = self.doctor_tree.selection()
        if selected:
            values = self.doctor_tree.item(selected[0])['values']
            for (field, value) in zip(self.doctor_entries.keys(), values):
                if isinstance(self.doctor_entries[field], ttk.Combobox):
                    self.doctor_entries[field].set(value)
                else:
                    self.doctor_entries[field].delete(0, tk.END)
                    self.doctor_entries[field].insert(0, value)

    def __del__(self):
        if self.conn:
            try:
                self.conn.commit()
                self.conn.close()
            except Exception as e:
                print(f"数据库关闭错误: {str(e)}")
            finally:
                self.conn = None
        #诊断系统模块，系统功能显示与判断模块
    def open_diagnostic_terminal(self):
        diagnostic_window = tk.Toplevel(self.root)
        diagnostic_window.title("智能诊断终端")
        diagnostic_window.geometry("1200x900")
        
        # 主框架
        main_frame = ttk.Frame(diagnostic_window, padding=20)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
       # 患者信息区域 - 减小高度
        info_frame = ttk.LabelFrame(main_frame, text="患者基本信息", padding=10)
        info_frame.pack(fill=tk.X, pady=5)  # 减小上下边距
        
       # 患者选择区域
        select_frame = ttk.Frame(info_frame)
        select_frame.pack(fill=tk.X, pady=2)  # 减小内部边距
        
        tk.Label(select_frame, text="选择患者:").pack(side=tk.LEFT, padx=5)
        self.diagnostic_patient_combo = ttk.Combobox(select_frame, width=30)
        self.diagnostic_patient_combo.pack(side=tk.LEFT, padx=5)
        ttk.Button(select_frame, text="自动诊断", 
                command=self.perform_diagnosis).pack(side=tk.LEFT, padx=5)
        ttk.Button(select_frame, text="保存诊断结果", 
                command=self.save_diagnosis_result).pack(side=tk.LEFT, padx=5)

        # 显示选中患者的基本信息
        self.patient_info_label = ttk.Label(info_frame, text="")
        self.patient_info_label.pack(pady=2)
        
        # 多模态数据汇总区域
        data_frame = ttk.LabelFrame(main_frame, text="多模态数据汇总", padding=10)
        data_frame.pack(fill=tk.BOTH, expand=True, pady=5)
        
        # 创建表格显示检查结果
        columns = ("检查项目", "检查结果", "参考结果", "检查时间")
        self.diagnostic_tree = ttk.Treeview(data_frame, columns=columns, show="headings", height=5)
        
        for col in columns:
            self.diagnostic_tree.heading(col, text=col)
            self.diagnostic_tree.column(col, width=150)
        
        scrollbar = ttk.Scrollbar(data_frame, orient=tk.VERTICAL, command=self.diagnostic_tree.yview)
        self.diagnostic_tree.configure(yscrollcommand=scrollbar.set)
        
        self.diagnostic_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # 诊断结果区域 - 增加高度
        result_frame = ttk.LabelFrame(main_frame, text="诊断结果", padding=15)
        result_frame.pack(fill=tk.BOTH, expand=True, pady=10)
        
        # 诊断结果文本框 - 增加高度
        self.diagnosis_result_text = tk.Text(result_frame, height=50, width=50)  # 增加height值
        self.diagnosis_result_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # 加载患者列表
        self.load_diagnostic_patients()
        
        # 绑定选择事件
        self.diagnostic_patient_combo.bind('<<ComboboxSelected>>', self.on_patient_selected_for_diagnosis)
    
    def load_diagnostic_patients(self):
        """加载患者列表"""
        try:
            self.cursor.execute('''
                SELECT patient_id, name, gender, age FROM patients
                ORDER BY admission_date DESC
            ''')
            patients = self.cursor.fetchall()
            self.diagnostic_patient_combo['values'] = [
                f"{p[0]} - {p[1]} ({p[2]}, {p[3]}岁)" for p in patients
            ]
        except Exception as e:
            messagebox.showerror("错误", f"加载患者列表失败：{str(e)}")
    
    def on_patient_selected_for_diagnosis(self, event=None):
        """当选择患者时更新信息显示"""
        if not self.diagnostic_patient_combo.get():
            return
            
        patient_id = self.diagnostic_patient_combo.get().split(" - ")[0]
        
        try:
            # 获取患者基本信息
            self.cursor.execute('''
                SELECT p.*, d.name as doctor_name
                FROM patients p
                LEFT JOIN doctors d ON p.doctor_id = d.doctor_id
                WHERE p.patient_id = ?
            ''', (patient_id,))
            
            patient = self.cursor.fetchone()
            if patient:
                info_text = f"病历号: {patient[0]} | 姓名: {patient[1]} | "
                info_text += f"性别: {patient[2]} | 年龄: {patient[3]} | "
                info_text += f"主治医生: {patient[-1] or '未指定'}"
                self.patient_info_label.config(text=info_text)
            
            # 清空现有数据
            self.diagnostic_tree.delete(*self.diagnostic_tree.get_children())
            
            # 获取心率血氧数据
            self.cursor.execute('''
                SELECT '心率检测（生理）', heart_rate || ' 次/分钟', 
                       '60-100次/分钟', measure_time
                FROM vital_signs 
                WHERE patient_id = ? 
                ORDER BY measure_time DESC LIMIT 1
            ''', (patient_id,))
            vital_sign = self.cursor.fetchone()
            if vital_sign:
                self.diagnostic_tree.insert("", tk.END, values=vital_sign)
                
            self.cursor.execute('''
                SELECT '血氧饱和度检测（生理）', blood_oxygen || ' %',
                       '95-100%', measure_time
                FROM vital_signs 
                WHERE patient_id = ?
                ORDER BY measure_time DESC LIMIT 1
            ''', (patient_id,))
            oxygen = self.cursor.fetchone()
            if oxygen:
                self.diagnostic_tree.insert("", tk.END, values=oxygen)
            
            # 获取医学影像分析结果
            self.cursor.execute('''
                SELECT '双心室核磁共振影像扫描（形态学）', morphology_result,
                       '形态学正常', analysis_date
                FROM medical_image 
                WHERE patient_id = ?
                ORDER BY analysis_date DESC LIMIT 1
            ''', (patient_id,))
            image = self.cursor.fetchone()
            if image:
                self.diagnostic_tree.insert("", tk.END, values=image)
            
                      # 获取电固液耦合分析结果
            self.cursor.execute('''
                SELECT '心室外心肌层的电势分布（电固耦合）', electrical_result, '正常范围', analysis_time
                FROM coupling_analysis 
                WHERE patient_id = ?
                ORDER BY analysis_time DESC LIMIT 1
            ''', (patient_id,))
            electrical = self.cursor.fetchone()
            if electrical:
                self.diagnostic_tree.insert("", tk.END, values=electrical)
            
            self.cursor.execute('''
                SELECT '心室心肌壁的应力分布（电固耦合）', stress_result, '正常范围', analysis_time
                FROM coupling_analysis 
                WHERE patient_id = ?
                ORDER BY analysis_time DESC LIMIT 1
            ''', (patient_id,))
            stress = self.cursor.fetchone()
            if stress:
                self.diagnostic_tree.insert("", tk.END, values=stress)
            
            self.cursor.execute('''
                SELECT '心室心肌壁的应变分布（电固耦合）', strain_result, '正常范围', analysis_time
                FROM coupling_analysis 
                WHERE patient_id = ?
                ORDER BY analysis_time DESC LIMIT 1
            ''', (patient_id,))
            strain = self.cursor.fetchone()
            if strain:
                self.diagnostic_tree.insert("", tk.END, values=strain)
            
            self.cursor.execute('''
                SELECT '心室内血流的流速分布（流体力学）', velocity_result, '正常范围', analysis_time
                FROM coupling_analysis 
                WHERE patient_id = ?
                ORDER BY analysis_time DESC LIMIT 1
            ''', (patient_id,))
            velocity = self.cursor.fetchone()
            if velocity:
                self.diagnostic_tree.insert("", tk.END, values=velocity)
            
            self.cursor.execute('''
                SELECT '心室内血流压力分布（流体力学）', pressure_result, '正常范围', analysis_time
                FROM coupling_analysis 
                WHERE patient_id = ?
                ORDER BY analysis_time DESC LIMIT 1
            ''', (patient_id,))
            pressure = self.cursor.fetchone()
            if pressure:
                self.diagnostic_tree.insert("", tk.END, values=pressure)
            
            self.cursor.execute('''
                SELECT '孤子现象（数据驱动）', vorticity_result, '正常范围', analysis_time
                FROM coupling_analysis 
                WHERE patient_id = ?
                ORDER BY analysis_time DESC LIMIT 1
            ''', (patient_id,))
            vorticity = self.cursor.fetchone()
            if vorticity:
                self.diagnostic_tree.insert("", tk.END, values=vorticity)
           
                
        except Exception as e:
            messagebox.showerror("错误", f"加载诊断数据失败：{str(e)}")
    #显示诊断结果
    def perform_diagnosis(self):
        if not self.diagnostic_patient_combo.get():
            messagebox.showwarning("警告", "请先选择患者！")
            return
                
        try:
            # 获取所有检查结果
            results = {}
            for item in self.diagnostic_tree.get_children():
                values = self.diagnostic_tree.item(item)['values']
                results[values[0]] = values[1]
            
            # 构建诊断提示语
            prompt = "作为一名经验丰富的心脏专科医生，请根据以下检查结果给出诊断意见：\n\n"
            for test, result in results.items():
                prompt += f"{test}: {result}\n"
            try:
                # 使用 OpenAI SDK 调用 DeepSeek API
                from openai import OpenAI
                client = OpenAI(
                    api_key="sk-febe1af78a224ceaadc61c408b7419d8",
                    base_url="https://api.deepseek.com/v1"
                )
                response = client.chat.completions.create(
                    model="deepseek-chat",
                    messages=[
                        {"role": "system", "content": "你是一位经验丰富的心脏专科医生。"},
                        {"role": "user", "content": prompt}
                    ],
                    temperature=0.7,
                    max_tokens=2000,
                    stream=False
                )
                diagnosis_text = response.choices[0].message.content
                
            except Exception as e:
                print(f"API错误: {str(e)}")  # 添加错误信息打印
                # 如果 API 调用失败，使用本地诊断逻辑
                diagnosis_text = "诊断结果：\n\n"
                
                # 1. 生理指标分析
                if '心率检测（生理）' in results:
                    heart_rate = float(results['心率检测（生理）'].split()[0])
                    if heart_rate < 60:
                        diagnosis_text += "1. 生理指标异常：心动过缓（<60次/分），需要进一步评估心脏传导系统功能\n"
                    elif heart_rate > 100:
                        diagnosis_text += "1. 生理指标异常：心动过速（>100次/分），建议进一步检查是否存在心律失常\n"
                    else:
                        diagnosis_text += "1. 生理指标：心率在正常范围内（60-100次/分）\n"
                
                if '血氧饱和度检测（生理）' in results:
                    spo2 = float(results['血氧饱和度检测（生理）'].split()[0])
                    if spo2 < 95:
                        diagnosis_text += "2. 血氧饱和度偏低（<95%），提示可能存在心肺功能异常\n"
                    else:
                        diagnosis_text += "2. 血氧饱和度正常（≥95%）\n"
                
                # 2. 形态学分析
                if '双心室核磁共振影像扫描（形态学）' in results:
                    morphology = results['双心室核磁共振影像扫描（形态学）']
                    if "正常" in morphology:
                        diagnosis_text += "3. 心脏形态学检查未见明显异常\n"
                    else:
                        diagnosis_text += "3. 心脏形态学检查发现异常，具体表现为：" + morphology + "\n"
                
        # 3. 电固耦合分析
                coupling_issues = []
                if '心室外心肌层的电势分布（电固耦合）' in results and "出现异常" in results['心室外心肌层的电势分布（电固耦合）']:
                    coupling_issues.append("心脏外纤维电势分布异常")
                if '心室心肌壁的应力分布（电固耦合）' in results and "出现异常" in results['心室心肌壁的应力分布（电固耦合）']:
                    coupling_issues.append("心脏外纤维应力分布异常")
                if '心室心肌壁的应变分布（电固耦合）' in results and "出现异常" in results['心室心肌壁的应变分布（电固耦合）']:
                    coupling_issues.append("心脏外纤维应变分布异常")
                
                if not coupling_issues:  # 修改这里的逻辑判断
                    diagnosis_text += "4. 电固耦合分析各项指标在正常范围内\n"
                else:
                    diagnosis_text += f"4. 电固耦合分析发现异常：{', '.join(coupling_issues)}，提示可能存在心肌功能异常\n"
                
                # 4. 流体力学分析
                fluid_issues = []
                if '心室内血流的流速分布（流体力学）' in results and "出现异常" in results['心室内血流的流速分布（流体力学）']:
                    fluid_issues.append("血流速度异常")
                if '心室内血流压力分布（流体力学）' in results and "出现异常" in results['心室内血流压力分布（流体力学）']:
                    fluid_issues.append("血流压力异常")
                
                if not fluid_issues:  # 修改这里的逻辑判断
                    diagnosis_text += "5. 流体力学分析各项指标在正常范围内\n"
                else:
                    diagnosis_text += f"5. 流体力学分析发现异常：{', '.join(fluid_issues)}，提示可能存在心室功能异常\n"
                
                # 5. 数据驱动分析
                if '孤子现象（数据驱动）' in results:
                    if "出现异常" in results['孤子现象（数据驱动）']:
                        diagnosis_text += "6. 数据驱动分析发现异常孤子现象，提示可能心室内学业存在问题\n"
                    else:
                        diagnosis_text += "6. 数据驱动分析未发现异常孤子现象\n"
                # 综合建议
                diagnosis_text += "\n综合建议：\n"
                if any("出现异常" in result for result in results.values()):
                    diagnosis_text += "建议进行进一步的专科检查和随访，必要时考虑介入治疗。\n"
                else:
                    diagnosis_text += "各项指标均在正常范围内，建议定期复查。\n"

                # 输出诊断结果
                print(diagnosis_text)
                # 添加推理规则引擎
                def apply_diagnostic_rules(results):
                    diagnosis_rules = []
                    severity_level = 0
                    
                    # 规则1: 多指标联合判断 - 心功能不全
                    if ('心率检测（生理）' in results and float(results['心率检测（生理）'].split()[0]) > 100 and
                        '血氧饱和度检测（生理）' in results and float(results['血氧饱和度检测（生理）'].split()[0]) < 95):
                        diagnosis_rules.append("可能存在心功能不全，建议进行心脏超声检查")
                        severity_level += 2
                    
                    # 规则2: 电固耦合与流体力学关联分析
                    if ('心室心肌壁的应力分布（电固耦合）' in results and "出现异常" in results['心室心肌壁的应力分布（电固耦合）'] and
                        '心室内血流压力分布（流体力学）' in results and "出现异常" in results['心室内血流压力分布（流体力学）']):
                        diagnosis_rules.append("心室收缩功能可能存在异常，建议进行心脏功能评估")
                        severity_level += 2
                    
                    # 规则3: 形态学与功能关联
                    if ('双心室核磁共振影像扫描（形态学）' in results and "异常" in results['双心室核磁共振影像扫描（形态学）']):
                        if any("出现异常" in results.get(key, "") for key in ['心室心肌壁的应力分布（电固耦合）', '心室内血流压力分布（流体力学）']):
                            diagnosis_rules.append("心室结构与功能均异常，提示可能存在心肌病变")
                            severity_level += 3
                    
                    # 规则4: 孤子现象特殊分析
                    if '孤子现象（数据驱动）' in results and "出现异常" in results['孤子现象（数据驱动）']:
                        if severity_level >= 3:
                            diagnosis_rules.append("多项指标异常并出现孤子现象，建议立即进行进一步检查")
                            severity_level += 2
                    
                    return diagnosis_rules, severity_level

                # 应用推理规则
                diagnostic_rules, severity = apply_diagnostic_rules(results)
                #添加推理规则结果到诊断文本
                diagnosis_text += "\n\n您好，我作为一名经验丰富的心脏专科医生，我已经根据以上的检查结果和推理规则给出诊断意见：推理系统分析结果：\n"
                for rule in diagnostic_rules:
                    diagnosis_text += f"- {rule}\n"
                diagnosis_text += f"\n综合严重程度评分：{severity} (0-10)\n"
                
                # 输出诊断结果
                print(diagnosis_text)
                
                # 更新诊断结果显示
                self.diagnosis_result_text.delete("1.0", tk.END)
                self.diagnosis_result_text.insert("1.0", diagnosis_text)
                    
        except Exception as e:
                    messagebox.showerror("错误", f"诊断过程出错：{str(e)}")
    
    # 添加保存诊断结果的方法
    def save_diagnosis_result(self):
            if not self.diagnostic_patient_combo.get():
                messagebox.showwarning("警告", "请先选择患者！")
                return
            diagnosis_text = self.diagnosis_result_text.get("1.0", tk.END).strip()
            if not diagnosis_text:
                messagebox.showwarning("警告", "没有诊断结果可保存！")
                return
            try:
                patient_id = self.diagnostic_patient_combo.get().split(" - ")[0]
                
                # 更新患者表中的诊断信息
                self.cursor.execute('''
                    UPDATE patients 
                    SET diagnosis = ?
                    WHERE patient_id = ?
                ''', (diagnosis_text, patient_id))
                self.conn.commit()
                messagebox.showinfo("成功", "诊断结果已保存到病例信息！")
                
            except Exception as e:
                messagebox.showerror("错误", f"保存诊断结果失败：{str(e)}")
    
    #第七个模块(医生远程会诊中心)
    def open_remote_consultation(self):
        consultation_window = tk.Toplevel(self.root)
        consultation_window.title("远程会诊系统")
        consultation_window.geometry("1200x800")
        
        # 主框架
        main_frame = ttk.Frame(consultation_window)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
        
        # 患者选择区域
        select_frame = ttk.LabelFrame(main_frame, text="患者信息", padding=10)
        select_frame.pack(fill=tk.X, pady=(0, 10))
        
        # 创建水平布局框架
        button_frame = ttk.Frame(select_frame)
        button_frame.pack(fill=tk.X, padx=5, pady=5)
        
        # 患者选择组件
        ttk.Label(button_frame, text="选择患者:").pack(side=tk.LEFT, padx=5)
        self.consultation_patient_combo = ttk.Combobox(button_frame, width=30)
        self.consultation_patient_combo.pack(side=tk.LEFT, padx=5)
        
        # AI助手按钮
        ttk.Button(
            button_frame, 
            text="启动大模型", 
            command=self.start_ai_consultation,
            style="Custom.TButton"
        ).pack(side=tk.LEFT, padx=20)
        
        # 诊断结果显示区域
        diagnosis_frame = ttk.LabelFrame(main_frame, text="诊断结果", padding=10)
        diagnosis_frame.pack(fill=tk.X, pady=(0, 10))

        # 创建滚动条
        diagnosis_scrollbar = ttk.Scrollbar(diagnosis_frame)
        diagnosis_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        # 修改诊断结果文本框并添加滚动条
        self.diagnosis_text = tk.Text(diagnosis_frame, height=4, width=80)  # 减小height值
        self.diagnosis_text.pack(fill=tk.X, padx=5, pady=5)

        # 配置滚动条
        self.diagnosis_text.config(yscrollcommand=diagnosis_scrollbar.set)
        diagnosis_scrollbar.config(command=self.diagnosis_text.yview)
        
        # 视频显示区域
        video_frame = ttk.LabelFrame(main_frame, text="视频会议", padding=10)
        video_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 10))
        
        # 视频显示的水平框架
        video_display_frame = ttk.Frame(video_frame)
        video_display_frame.pack(fill=tk.BOTH, expand=True)
        
        # 本地视频显示
        local_frame = ttk.LabelFrame(video_display_frame, text="本地视频", padding=5)
        local_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 5))
        self.local_canvas = tk.Canvas(local_frame, width=380, height=250, bg='black')
        self.local_canvas.pack(pady=5)
        
        # 远程视频显示
        remote_frame = ttk.LabelFrame(video_display_frame, text="远程视频", padding=5)
        remote_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(5, 0))
        
        self.remote_canvas = tk.Canvas(remote_frame, width=380, height=250, bg='black')
        self.remote_canvas.pack(pady=5)
        
        # 控制按钮
        control_frame = ttk.Frame(main_frame)
        control_frame.pack(fill=tk.X, pady=(0, 10))
        
        self.start_btn = ttk.Button(control_frame, text="开始会诊", 
                                command=self.start_video,
                                style="Custom.TButton")
        self.start_btn.pack(side=tk.LEFT, padx=5)
        
        self.stop_btn = ttk.Button(control_frame, text="结束会诊", 
                                command=self.stop_video,
                                state='disabled',
                                style="Custom.TButton")
        self.stop_btn.pack(side=tk.LEFT, padx=5)
        
        # 加载患者列表
        self.load_consultation_patients()
        
        # 绑定选择事件
        self.consultation_patient_combo.bind('<<ComboboxSelected>>', 
                                        self.on_patient_selected_for_consultation)
        
        self.cap = None
        self.is_running = False

    def load_consultation_patients(self):
        try:
            self.cursor.execute('''
                SELECT patient_id, name, gender, age FROM patients
                ORDER BY admission_date DESC
            ''')
            patients = self.cursor.fetchall()
            self.consultation_patient_combo['values'] = [
                f"{p[0]} - {p[1]} ({p[2]}, {p[3]}岁)" for p in patients
            ]
        except Exception as e:
            messagebox.showerror("错误", f"加载患者列表失败：{str(e)}")

    def on_patient_selected_for_consultation(self, event=None):
        if not self.consultation_patient_combo.get():
            return
            
        patient_id = self.consultation_patient_combo.get().split(" - ")[0]
        
        try:
            # 获取患者诊断信息
            self.cursor.execute('''
                SELECT diagnosis FROM patients WHERE patient_id = ?
            ''', (patient_id,))
            
            diagnosis = self.cursor.fetchone()
            if diagnosis and diagnosis[0]:
                self.diagnosis_text.delete("1.0", tk.END)
                self.diagnosis_text.insert("1.0", diagnosis[0])
            else:
                self.diagnosis_text.delete("1.0", tk.END)
                self.diagnosis_text.insert("1.0", "暂无诊断信息")
                
        except Exception as e:
            messagebox.showerror("错误", f"获取诊断信息失败：{str(e)}")

    def start_video(self):
        try:
            # 初始化人脸识别和手势识别
            self.face_cascade = cv2.CascadeClassifier(cv2.data.haarcascades + 'haarcascade_frontalface_default.xml')
            self.mp_hands = mp.solutions.hands
            self.hands = self.mp_hands.Hands(max_num_hands=1, min_detection_confidence=0.7)
            self.mp_draw = mp.solutions.drawing_utils
            # 加载参考图片
            ref_image_path = r"C:\Users\32572\Desktop\认证资料与电子证书\30kb.jpg"
            ref_image = cv2.imread(ref_image_path)
            if ref_image is None:
                messagebox.showerror("错误", "无法加载参考图片，请检查图片路径是否正确")
                return
            ref_gray = cv2.cvtColor(ref_image, cv2.COLOR_BGR2GRAY)
            
            # 检测参考图片中的人脸，调整参数使其更容易检测到人脸
            ref_faces = self.face_cascade.detectMultiScale(ref_gray, scaleFactor=1.1, minNeighbors=3)
            if len(ref_faces) > 0:
                x, y, w, h = ref_faces[0]
                self.ref_face = ref_gray[y:y + h, x:x + w]
                self.orb = cv2.ORB_create(nfeatures=1000)  # 增加特征点数量
                self.kp_ref, self.des_ref = self.orb.detectAndCompute(self.ref_face, None)
                self.bf = cv2.BFMatcher(cv2.NORM_HAMMING, crossCheck=True)
            else:
                messagebox.showerror("错误", "参考图片中未检测到人脸")
                return
            
            # 初始化摄像头
            for camera_index in range(2):
                self.cap = cv2.VideoCapture(camera_index)
                if self.cap.isOpened():
                    break
            if not self.cap.isOpened():
                messagebox.showerror("错误", "无法打开摄像头！请检查：\n1. 摄像头是否正确连接\n2. 是否被其他程序占用\n3. Windows相机隐私设置是否允许访问")
                return
            
            self.cap.set(cv2.CAP_PROP_FRAME_WIDTH, 640)
            self.cap.set(cv2.CAP_PROP_FRAME_HEIGHT, 480)
            self.is_running = True
            self.authenticated = False
            self.auth_show_counter = 0
            self.start_btn.config(state='disabled')
            self.stop_btn.config(state='normal')
            self.update_video()
        except Exception as e:
            messagebox.showerror("错误", f"启动摄像头失败：{str(e)}")
            print(f"Camera error details: {str(e)}")
    def update_video(self):
        if self.is_running and self.cap is not None:
            try:
                ret, frame = self.cap.read()
                if ret:
                    frame_copy = frame.copy()
                    gray = cv2.cvtColor(frame_copy, cv2.COLOR_BGR2GRAY)
                    faces = self.face_cascade.detectMultiScale(
                        gray,
                        scaleFactor=1.1,
                        minNeighbors=3,
                        minSize=(30, 30)
                    )
                    # 无论是否认证都显示人脸框
                    for (x, y, w, h) in faces:
                        cv2.rectangle(frame_copy, (x, y), (x+w, y+h), (255, 0, 0), 2)
                        if not self.authenticated:
                            face_roi = gray[y:y + h, x:x + w]
                            kp_frame, des_frame = self.orb.detectAndCompute(face_roi, None)
                            
                            if des_frame is not None and self.des_ref is not None:
                                matches = self.bf.match(self.des_ref, des_frame)
                                matches = sorted(matches, key=lambda x: x.distance)
                                
                                if len(matches) > 5:
                                    self.authenticated = True
                    
                    # 显示认证状态
                    if not self.authenticated:
                        cv2.putText(frame_copy, "Authenticating...", (10, 30), 
                                cv2.FONT_HERSHEY_SIMPLEX, 1, (0, 0, 255), 2)
                    else:
                        cv2.putText(frame_copy, "Authenticated", (10, 30), 
                                cv2.FONT_HERSHEY_SIMPLEX, 1, (0, 255, 0), 2)
                        
                        # 认证后进行手势识别
                        rgb_frame = cv2.cvtColor(frame_copy, cv2.COLOR_BGR2RGB)
                        results = self.hands.process(rgb_frame)
                        
                        if results.multi_hand_landmarks:
                            for hand_landmarks in results.multi_hand_landmarks:
                                self.mp_draw.draw_landmarks(frame_copy, hand_landmarks, 
                                                        self.mp_hands.HAND_CONNECTIONS)
                    
                    # 转换并显示图像
                    frame_copy = cv2.flip(frame_copy, 1)  # 水平翻转
                    frame_copy = cv2.cvtColor(frame_copy, cv2.COLOR_BGR2RGB)
                    frame_copy = cv2.resize(frame_copy, (400, 300))
                    
                    image = Image.fromarray(frame_copy)
                    photo = ImageTk.PhotoImage(image=image)
                    
                    self.local_canvas.create_image(0, 0, image=photo, anchor=tk.NW)
                    self.local_canvas.photo = photo
                
                if self.is_running:
                    self.root.after(33, self.update_video)
                    
            except Exception as e:
                print(f"Frame update error: {str(e)}")
                self.stop_video()                    
    
    def start_ai_consultation(self):
        try:
            import subprocess
            import webbrowser
            import os
            
            # 获取chat_app.py的完整路径
            app_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'untitled1.py')
            
            # 使用subprocess启动web服务器
            subprocess.Popen(['python', '-m', 'streamlit', 'run', app_path])
            
            # 等待几秒钟让服务器启动
            import time
            time.sleep(3)
            
            # 自动打开浏览器
            webbrowser.open('http://localhost:8501')
            
            messagebox.showinfo("成功", "AI助手已启动！请在浏览器中使用。")
        except Exception as e:
            messagebox.showerror("错误", f"启动AI助手失败：{str(e)}")
    def stop_video(self):
        self.is_running = False
        if self.cap:
            self.cap.release()
        self.start_btn.config(state='normal')
        self.stop_btn.config(state='disabled')
#血液科项目检验
    def open_biochemical_test(self):
        test_window = tk.Toplevel(self.root)
        test_window.title("生化检验全科")
        test_window.geometry("1200x800")
        # 主框架
        main_frame = ttk.Frame(test_window, padding=20)
        main_frame.pack(fill=tk.BOTH, expand=True)
        # 患者选择区域
        self.select_frame = ttk.LabelFrame(main_frame, text="患者信息", padding=10)
        self.select_frame.pack(fill=tk.X, pady=10)
        ttk.Label(self.select_frame, text="选择患者:").pack(side=tk.LEFT, padx=5)
        self.test_patient_combo = ttk.Combobox(self.select_frame, width=30)
        self.test_patient_combo.pack(side=tk.LEFT, padx=5)
        # 检验项目区域 - 使用滚动画布
        test_canvas = tk.Canvas(main_frame)
        scrollbar = ttk.Scrollbar(main_frame, orient="vertical", command=test_canvas.yview)
        test_frame = ttk.LabelFrame(test_canvas, text="生化指标", padding=10)
        test_canvas.configure(yscrollcommand=scrollbar.set)
        # 创建内部框架
        inner_frame = ttk.Frame(test_frame)
        inner_frame.pack(fill=tk.BOTH, expand=True)
        # 创建检验项目输入框
        self.test_entries = {}
        
        # 分组定义检验项目
        cardiac_markers = [
            ("肌钙蛋白I(cTnI)", "ng/mL", "0-0.04"),
            ("肌酸激酶同工酶(CK-MB)", "U/L", "0-25"),
            ("脑钠肽(BNP)", "pg/mL", "0-100"),
            ("乳酸脱氢酶(LDH)", "U/L", "120-250"),
            ("天门冬氨酸氨基转移酶(AST)", "U/L", "15-40"),
            ("肌酸激酶(CK)", "U/L", "50-310"),
            ("肌红蛋白(Mb)", "ng/mL", "0-110"),
            ("C反应蛋白(CRP)", "mg/L", "0-8")
        ]
        
        

        # 创建表头
        headers = ["检验项目", "检测值", "单位", "参考范围", "结果判断"]
        for col, header in enumerate(headers):
            ttk.Label(inner_frame, text=header, font=("Microsoft YaHei", 10, "bold")).grid(
                row=0, column=col, padx=5, pady=5)
        
        current_row = 1
        
        # 心脏标志物
        ttk.Label(inner_frame, text="心脏病标志物检验指标", font=("Microsoft YaHei", 10, "bold")).grid(
            row=current_row, column=0, columnspan=5, padx=5, pady=10, sticky="w")
        current_row += 1
        
        for item, unit, ref_range in cardiac_markers:
            self.create_test_item(inner_frame, item, unit, ref_range, current_row, 0)
            current_row += 1
       
        # 配置滚动区域
        test_frame.pack(fill=tk.BOTH, expand=True)
        test_canvas.create_window((0, 0), window=test_frame, anchor="nw")
        
        # 更新滚动区域
        test_frame.update_idletasks()
        test_canvas.configure(scrollregion=test_canvas.bbox("all"))
        
        # 放置画布和滚动条
        test_canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, pady=10)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # 控制按钮区域
        control_frame = ttk.Frame(main_frame)
        control_frame.pack(fill=tk.X, pady=10)
        
        ttk.Button(control_frame, text="保存结果", 
                  command=self.save_biochemical_results).pack(side=tk.LEFT, padx=5)
        ttk.Button(control_frame, text="生成报告", 
                  command=self.generate_biochemical_report).pack(side=tk.LEFT, padx=5)
        
        # 加载患者列表
        self.load_biochemical_patients()
    
    def create_test_item(self, parent, item, unit, ref_range, row, col_offset):
        # 项目名称
        ttk.Label(parent, text=item).grid(
            row=row, column=col_offset, padx=5, pady=5)
        
        # 检测值输入框
        value_entry = ttk.Entry(parent, width=10)
        value_entry.grid(row=row, column=col_offset+1, padx=5, pady=5)
        self.test_entries[item] = value_entry
        
        # 单位
        ttk.Label(parent, text=unit).grid(
            row=row, column=col_offset+2, padx=5, pady=5)
        
        # 参考范围
        ttk.Label(parent, text=ref_range).grid(
            row=row, column=col_offset+3, padx=5, pady=5)
        
        # 结果判断
        result_combo = ttk.Combobox(parent, values=["正常", "偏高", "偏低"], 
                                  width=8, state="readonly")
        result_combo.grid(row=row, column=col_offset+4, padx=5, pady=5)
        result_combo.set("正常")    
    def load_biochemical_patients(self):
        try:
            self.cursor.execute('''
                SELECT p.patient_id, p.name, p.gender, p.age, d.name as doctor_name 
                FROM patients p
                LEFT JOIN doctors d ON p.doctor_id = d.doctor_id
            ''')
            patients = self.cursor.fetchall()
            self.test_patient_combo['values'] = [f"{p[0]} - {p[1]}" for p in patients]
            
            # 绑定选择事件
            self.test_patient_combo.bind('<<ComboboxSelected>>', self.on_biochemical_patient_selected)
        except Exception as e:
            messagebox.showerror("错误", f"加载患者列表失败：{str(e)}")
    def on_biochemical_patient_selected(self, event=None):
        if not self.test_patient_combo.get():
            return
            
        patient_id = self.test_patient_combo.get().split(" - ")[0]
        
        try:
            # 获取患者详细信息
            self.cursor.execute('''
                SELECT p.*, d.name as doctor_name
                FROM patients p
                LEFT JOIN doctors d ON p.doctor_id = d.doctor_id
                WHERE p.patient_id = ?
            ''', (patient_id,))
            
            patient = self.cursor.fetchone()
            if patient:
                # 在患者信息框中显示信息
                info_text = f"病历号: {patient[0]} | 姓名: {patient[1]} | "
                info_text += f"性别: {patient[2]} | 年龄: {patient[3]} | "
                info_text += f"主治医生: {patient[-1] or '未指定'}"
                
                if hasattr(self, 'patient_info_label'):
                    self.patient_info_label.config(text=info_text)
                else:
                    self.patient_info_label = ttk.Label(self.select_frame, text=info_text)  # 使用实例变量
                    self.patient_info_label.pack(side=tk.LEFT, padx=5)
                
        except Exception as e:
            messagebox.showerror("错误", f"获取患者信息失败：{str(e)}")        

    def save_biochemical_results(self):
        if not self.test_patient_combo.get():
            messagebox.showwarning("警告", "请先选择患者！")
            return
            
        try:
            patient_id = self.test_patient_combo.get().split(" - ")[0]
            
            # 只收集心脏病相关的检验值
            cardiac_test_values = {}
            cardiac_markers = [
                "肌钙蛋白I(cTnI)",
                "肌酸激酶同工酶(CK-MB)",
                "脑钠肽(BNP)",
                "乳酸脱氢酶(LDH)",
                "天门冬氨酸氨基转移酶(AST)",
                "肌酸激酶(CK)",
                "肌红蛋白(Mb)",
                "C反应蛋白(CRP)"
            ]
            
            for marker in cardiac_markers:
                if marker in self.test_entries:
                    value = self.test_entries[marker].get()
                    if value:
                        cardiac_test_values[marker.split('(')[0]] = float(value)
                    
            # 插入数据库
            self.cursor.execute('''
                INSERT INTO biochemical_tests 
                (patient_id, cTnI, CK_MB, BNP, LDH, AST, CK, Mb, CRP)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (
                patient_id,
                cardiac_test_values.get('肌钙蛋白I', None),
                cardiac_test_values.get('肌酸激酶同工酶', None),
                cardiac_test_values.get('脑钠肽', None),
                cardiac_test_values.get('乳酸脱氢酶', None),
                cardiac_test_values.get('天门冬氨酸氨基转移酶', None),
                cardiac_test_values.get('肌酸激酶', None),
                cardiac_test_values.get('肌红蛋白', None),
                cardiac_test_values.get('C反应蛋白', None)
            ))
            
            self.conn.commit()
            messagebox.showinfo("成功", "心脏病相关检验结果已保存！")
            
        except Exception as e:
            messagebox.showerror("错误", f"保存结果失败：{str(e)}")

    def generate_biochemical_report(self):
        if not self.test_patient_combo.get():
            messagebox.showwarning("警告", "请先选择患者！")
            return
        try:
            patient_id = self.test_patient_combo.get().split(" - ")[0]
            
            # 获取患者信息
            self.cursor.execute('''
                SELECT p.*, d.name as doctor_name
                FROM patients p
                LEFT JOIN doctors d ON p.doctor_id = d.doctor_id
                WHERE p.patient_id = ?
            ''', (patient_id,))
            patient = self.cursor.fetchone()
            # 准备CSV数据
            import csv
            from datetime import datetime
            
            # 生成文件名
            current_time = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"c:\\Users\\32572\\Desktop\\检验报告_{patient_id}_{current_time}.csv"
            
            # 写入CSV文件
            with open(filename, 'w', newline='', encoding='utf-8-sig') as f:
                writer = csv.writer(f)
                
                # 写入报告头部信息
                writer.writerow(['检验报告'])
                writer.writerow(['报告时间', datetime.now().strftime("%Y-%m-%d %H:%M:%S")])
                writer.writerow(['患者信息'])
                writer.writerow(['病历号', patient[0]])
                writer.writerow(['姓名', patient[1]])
                writer.writerow(['性别', patient[2]])
                writer.writerow(['年龄', patient[3]])
                writer.writerow(['主治医生', patient[-1] or '未指定'])
                writer.writerow([])  # 空行
                
                # 写入检验结果标题
                writer.writerow(['检验项目', '检测值', '单位', '参考范围', '结果判断'])
                
                # 定义所有检验项目的单位和参考范围
                test_info = {
                    # 心脏标志物
                    "肌钙蛋白I(cTnI)": ("ng/mL", "0-0.04"),
                    "肌酸激酶同工酶(CK-MB)": ("U/L", "0-25"),
                    "脑钠肽(BNP)": ("pg/mL", "0-100"),
                    # ... 其他检验项目信息 ...
                }
                # 获取所有已填写的检验结果
                for item, entry in self.test_entries.items():
                    value = entry.get()
                    if value:  # 只保存已填写的结果
                        # 获取该项目所在行的其他信息
                        parent = entry.master
                        row_info = entry.grid_info()['row']
                        
                        # 从预定义信息中获取单位和参考范围
                        unit, ref_range = test_info.get(item, ("", ""))
                        
                        # 获取结果判断
                        result_widgets = [w for w in parent.grid_slaves(row=row_info) if isinstance(w, ttk.Combobox)]
                        judgment = result_widgets[0].get() if result_widgets else "未判断"
                        
                        # 写入该行数据
                        writer.writerow([item, value, unit, ref_range, judgment])
            
            messagebox.showinfo("成功", f"报告已保存为CSV文件：\n{filename}")
        except Exception as e:
            messagebox.showerror("错误", f"生成报告失败：{str(e)}")
if __name__ == "__main__":
    LoginWindow()